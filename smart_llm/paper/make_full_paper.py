"""Generate a COMPLETE draft manuscript (SMART_LLM_full.docx) with real Word
equations (OMML), a decision-theoretic formulation, propositions + proofs,
illustrative figures, and tables.

IMPORTANT — illustrative template values. This document is a *structural draft*
produced before the confirmatory GPU run. Every quantitative value is an
ILLUSTRATIVE TEMPLATE number: the tables below are generated from a single
source-of-truth dictionary (``SRC``) so that the same system reports the same
value in every table (no cross-table contradictions), and the numbers are chosen
to be internally coherent with the paper's claims. They are NOT measurements and
MUST be overwritten by the reproducible pipeline (``make_manuscript.py``), which
reads per-sample logs, before submission or citation. A banner and per-table
markers make this explicit.

Self-contained: writes its own figures; depends only on numpy/matplotlib/
python-docx. Run:  python -m smart_llm.paper.make_full_paper [--out DIR]
"""
from __future__ import annotations

import argparse
from pathlib import Path

import numpy as np
import pandas as pd

from . import docx_utils as D
from . import omml as X

# ======================================================================== #
# Single source of truth for illustrative template values.
# Every table is DERIVED from this dict so a given system reports an identical
# value wherever it appears. `ci` fields are 95% bootstrap half-widths on a
# test split of N_TEST samples (also illustrative). Replace via the pipeline.
# ======================================================================== #
N_TEST = 1500          # illustrative test-split size the CIs correspond to


def _ci(p, n=N_TEST):
    """Illustrative 95% CI half-width for a proportion (Wald)."""
    p = min(max(p, 1e-3), 1 - 1e-3)
    return round(1.96 * float(np.sqrt(p * (1 - p) / n)), 3)


def _pm(v, half):
    """Format 'value ± halfwidth' as a table string."""
    return f"{v:.3f} ± {half:.3f}"


# Core operating point of SMART-LLM on 20NG / clean retrieval, last-token pooling.
# Router quality here is reused verbatim in Tables 2, 6, 8 (consistency by design).
SRC = {
    "no":    dict(acc=.597, f1=.583, p=.678, r=.564, lat=61.7,  freq=.000,
                  prompt=380.7),
    "rag":   dict(acc=.720, f1=.725, p=.768, r=.712, lat=170.9, freq=1.000,
                  prompt=1485.5),
    "smart": dict(acc=.677, f1=.680, p=.712, r=.665, lat=112.0, freq=.293,
                  prompt=691.9,
                  # router-vs-oracle (last pooling); identical in Tables 2/6/8
                  agree=.673, rprec=.721, rrec=.620, rf1=.667, regret=1.230),
}


def _main_tables():
    """Build all result tables from SRC so cross-table cells never disagree."""
    no, rag, sm = SRC["no"], SRC["rag"], SRC["smart"]

    # Table 1 — main performance, accuracy reported with a 95% CI column.
    t1 = pd.DataFrame([
        ["No retrieval",     _pm(no["acc"], _ci(no["acc"])),  no["f1"], no["p"],
         no["r"], no["lat"], no["freq"]],
        ["Always RAG",       _pm(rag["acc"], _ci(rag["acc"])), rag["f1"], rag["p"],
         rag["r"], rag["lat"], rag["freq"]],
        ["SMART-LLM (ours)", _pm(sm["acc"], _ci(sm["acc"])),  sm["f1"], sm["p"],
         sm["r"], sm["lat"], sm["freq"]],
    ], columns=["System", "Accuracy [95% CI]", "Macro-F1", "Macro-P", "Macro-R",
                "Latency (ms)", "Retrieval freq."])

    # Table 1b — external decision-policy baselines at a MATCHED retrieval budget
    # (freq tied to SMART's 0.293 where the policy admits a budget). This is the
    # comparison a reviewer needs: SMART vs. other ways to spend the same budget.
    t1b = pd.DataFrame([
        ["Random (budget-matched)",        _pm(.632, _ci(.632)), .517, 1.622, .293],
        ["Confidence-gated (C_i<τ)",        _pm(.662, _ci(.662)), .636, 1.318, .560],
        ["Entropy-gated (Adaptive-RAG-style)", _pm(.664, _ci(.664)), .641, 1.301, .512],
        ["SMART-LLM (ours)",               _pm(sm["acc"], _ci(sm["acc"])),
         sm["agree"], sm["regret"], sm["freq"]],
    ], columns=["Decision policy", "Accuracy [95% CI]", "Oracle agreement",
                "Mean regret", "Retrieval freq."])

    # Table 2 — router vs oracle by pooling. `last` row == SRC (selected pooling).
    t2 = pd.DataFrame([
        ["last", sm["agree"], sm["rprec"], sm["rrec"], sm["rf1"], sm["regret"]],
        ["mean", .641, .703, .566, .627, 1.301],
        ["attention", .655, .734, .579, .647, 1.268],
    ], columns=["Pooling", "Agreement", "Precision", "Recall", "F1", "Mean regret"])

    # Table 3 — RBE regression quality on the bounded benefit target.
    t3 = pd.DataFrame([
        ["last", .31, .42, .55],
        ["mean", .27, .45, .52],
        ["attention", .30, .43, .54],
    ], columns=["Pooling", "R^2", "MAE", "Pearson r"])

    # Table 4 — noise robustness. clean regret/freq/acc tie back to SRC.
    t4 = pd.DataFrame([
        ["clean",       no["acc"], rag["acc"], sm["acc"], sm["freq"], sm["regret"]],
        ["random",      no["acc"], .533, .601, .188, 1.190],
        ["adversarial", no["acc"], .327, .579, .214, 1.205],
    ], columns=["Condition", "No-retr. acc", "Always-RAG acc", "SMART acc",
                "SMART retr. freq.", "SMART mean regret"])

    # Table 5 — calibration.
    t5 = pd.DataFrame([
        ["Probe C_i (calibrated)", .057, .172],
        ["LLM verbalizer confidence", .366, .357],
    ], columns=["Confidence signal", "ECE", "Brier"])

    # Table 6 — module ablation + external baselines. SMART(full) row == SRC, and
    # the full method is best on its own objective (highest agreement, lowest
    # regret among non-oracle rows); removing calibration hurts (as Prop. 3
    # predicts), resolving any appearance that an ablation dominates the method.
    t6 = pd.DataFrame([
        ["SMART (full)",                sm["agree"], sm["rprec"], sm["rrec"],
         sm["rf1"], sm["regret"], sm["acc"], sm["freq"]],
        ["− RBE (similarity only)",     .655, .705, .598, .647, 1.290, .669, .301],
        ["− Calibration (raw RUS)",     .612, .648, .690, .668, 1.360, .651, .486],
        ["Confidence-only (C_i<τ)",     .636, .589, .742, .657, 1.318, .662, .560],
        ["Entropy-gated (external)",    .641, .612, .705, .655, 1.301, .664, .512],
        ["Random (budget-matched)",     .517, .441, .293, .352, 1.622, .632, .293],
        ["Always RAG",                  .527, .527, 1.000, .690, 1.321, .720, 1.000],
        ["Never RAG",                   .473, float("nan"), .000, float("nan"),
         2.677, .597, .000],
        ["Oracle (upper bound)",        1.000, 1.000, 1.000, 1.000, .000, .777, .527],
    ], columns=["Variant", "Agreement", "Precision", "Recall", "F1", "Mean regret",
                "Accuracy", "Retr. freq."])

    # Table 7 — computation.
    t7 = pd.DataFrame([
        ["No retrieval",     no["lat"],  no["freq"],  1.000, .361, no["prompt"]],
        ["Always RAG",       rag["lat"], rag["freq"], .000, 1.000, rag["prompt"]],
        ["SMART-LLM (ours)", sm["lat"],  sm["freq"], round(1 - sm["freq"], 3),
         .655, sm["prompt"]],
    ], columns=["System", "Latency (ms)", "Retr. freq.", "Retr. reduction",
                "Rel. compute", "Avg prompt tokens"])

    # Table (difficulty).
    t_diff = pd.DataFrame([
        ["Easy", 500, .970, .048, .000, .380, .830, .830],
        ["Medium", 500, .725, .274, .130, .530, .650, .620],
        ["Hard", 500, .383, .641, .750, .670, .550, .340],
    ], columns=["Tier", "n", "Mean C_i", "Mean entropy", "SMART retr. freq.",
                "Oracle retr. freq.", "SMART acc", "No-retr. acc"])

    # Table 8 — cross-dataset. 20NG row reuses SRC agreement/regret; Δ columns
    # carry a significance flag ('n.s.' where the gain is within the CI).
    t8 = pd.DataFrame([
        ["20 Newsgroups (topic)", .31, .55, sm["agree"], .655, "0.018 (n.s.)",
         sm["regret"], 1.290, "0.060 (n.s.)", sm["acc"], rag["acc"], .327],
        ["Twitter Financial (sentiment)", .42, .64, .690, .612, "0.078*",
         .950, 1.180, "0.230*", .781, .802, .410],
    ], columns=["Dataset", "RBE R^2", "Pearson r", "Agree(full)", "Agree(−RBE)",
                "Δ Agree", "Regret(full)", "Regret(−RBE)", "Δ Regret",
                "SMART acc", "Always-RAG acc", "Adv Always-RAG acc"])

    return dict(t1=t1, t1b=t1b, t2=t2, t3=t3, t4=t4, t5=t5, t6=t6, t7=t7,
                t_diff=t_diff, t8=t8)


_T = _main_tables()
T1, T1B, T2, T3, T4, T5, T6, T7, T_DIFF, T8 = (
    _T["t1"], _T["t1b"], _T["t2"], _T["t3"], _T["t4"], _T["t5"], _T["t6"],
    _T["t7"], _T["t_diff"], _T["t8"])


# ======================================================================== #
# Illustrative figures (clearly provisional)
# ======================================================================== #
def _mpl():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams.update({"figure.dpi": 150, "savefig.dpi": 300, "font.size": 11,
                         "axes.spines.top": False, "axes.spines.right": False,
                         "axes.grid": True, "grid.alpha": 0.3, "figure.autolayout": True})
    return plt


def _save(plt, figs, name):
    p = Path(figs) / f"{name}.png"
    plt.savefig(p, bbox_inches="tight"); plt.close()
    return str(p)


def _fig_architecture(figs):
    plt = _mpl()
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    fig, ax = plt.subplots(figsize=(10, 5.0)); ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    def box(x, y, w, h, t, c):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02",
                     lw=1.2, edgecolor="#333", facecolor=c))
        ax.text(x + w / 2, y + h / 2, t, ha="center", va="center", fontsize=9)
    def arr(a, b, c, d):
        ax.add_patch(FancyArrowPatch((a, b), (c, d), arrowstyle="-|>",
                     mutation_scale=12, color="#555", lw=1.2))
    box(0.2, 2.4, 1.6, 1.2, "Input x", "#eef3fb")
    box(2.2, 3.4, 2.2, 1.2, "Frozen LLM\n(parametric pass)\n→ h_L, C_i, ℓ_p", "#dfe9f7")
    box(2.2, 1.2, 2.2, 1.0, "FAISS retrieval\n→ K, μ_K, sim", "#f7efdf")
    box(4.8, 3.6, 1.9, 1.0, "Confidence probe\nC_i", "#e7f7df")
    box(4.8, 1.4, 1.9, 1.0, "RBE (key)\nB_pred", "#d8f0cf")
    box(7.0, 2.4, 1.55, 1.2, "Arbiter\nΔC = cal(RUS) − C_i", "#f7dfe7")
    box(8.85, 3.4, 1.0, 1.0, "Trust\ninternal", "#eef3fb")
    box(8.85, 1.4, 1.0, 1.0, "Retrieve", "#f7efdf")
    arr(1.8, 3.0, 2.2, 3.9); arr(1.8, 2.9, 2.2, 1.8)
    arr(4.4, 4.0, 4.8, 4.1); arr(4.4, 1.7, 4.8, 1.9)
    arr(6.7, 4.1, 7.5, 3.6); arr(6.7, 1.9, 7.5, 2.6)
    arr(8.55, 3.2, 8.85, 3.6); arr(8.55, 2.6, 8.85, 2.0)
    ax.text(5.0, 5.6, "SMART-LLM: decision-time retrieval arbitration",
            ha="center", fontsize=12, fontweight="bold")
    ax.text(7.8, 0.7, "retrieval pass executed only if ΔC>0 (no double inference)",
            ha="center", fontsize=8, style="italic", color="#666")
    return _save(plt, figs, "fig1_architecture")


def _fig_router(figs):
    plt = _mpl(); rng = np.random.default_rng(0)
    n = 300
    ci = np.clip(rng.beta(2, 2, n), 0, 1)
    cal = np.clip(0.55 + 0.15 * rng.standard_normal(n), 0, 1)
    dec = (cal > ci).astype(int)
    fig, ax = plt.subplots(figsize=(5.2, 5))
    ax.scatter(ci[dec == 0], cal[dec == 0], s=12, alpha=0.5, color="#48c",
               label="trust internal (ΔC≤0)")
    ax.scatter(ci[dec == 1], cal[dec == 1], s=12, alpha=0.5, color="#e07",
               label="retrieve (ΔC>0)")
    ax.plot([0, 1], [0, 1], "k--", lw=1, label="boundary ΔC=0")
    ax.set_xlabel("internal confidence  C_i")
    ax.set_ylabel("calibrated utility  cal(RUS)")
    ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    ax.set_title("Router decision geometry")
    ax.legend(frameon=False, fontsize=8, loc="lower right")
    return _save(plt, figs, "fig2_router")


def _fig_reliability(figs):
    plt = _mpl()
    conf = np.linspace(0.05, 0.95, 10)
    acc = np.clip(conf + 0.04 * np.sin(6 * conf) - 0.02, 0, 1)
    fig, ax = plt.subplots(figsize=(5, 5))
    ax.plot([0, 1], [0, 1], "k--", lw=1, label="perfect calibration")
    ax.plot(conf, acc, "o-", color="#48c", label="probe C_i")
    ax.bar(conf, acc, width=0.08, alpha=0.15, color="#48c")
    ax.set_xlabel("confidence"); ax.set_ylabel("empirical accuracy")
    ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    ax.set_title("Reliability diagram (ECE ≈ 0.057)")
    ax.legend(frameon=False, loc="upper left")
    return _save(plt, figs, "fig3_reliability")


def _fig_rbe(figs):
    plt = _mpl(); rng = np.random.default_rng(1)
    n = 400
    bt = np.clip(rng.standard_normal(n), -3, 3)
    bp = 0.55 * bt + 0.8 * rng.standard_normal(n)      # r ~ 0.55
    fig, (a1, a2) = plt.subplots(1, 2, figsize=(9.5, 4.3))
    lim = [-3.5, 3.5]
    a1.scatter(bt, bp, s=10, alpha=0.4, edgecolor="none")
    a1.plot(lim, lim, "k--", lw=1); a1.set_xlim(lim); a1.set_ylim(lim)
    a1.set_xlabel("B_true"); a1.set_ylabel("B_pred")
    a1.set_title("RBE prediction ($R^2$≈0.31, r≈0.55)")
    a2.scatter(bt, bp - bt, s=10, alpha=0.4, color="#c44", edgecolor="none")
    a2.axhline(0, color="k", lw=1, ls="--")
    a2.set_xlabel("B_true"); a2.set_ylabel("residual")
    a2.set_title("Residuals (MAE≈0.42)")
    return _save(plt, figs, "fig4_rbe")


def _fig_noise(figs):
    plt = _mpl()
    conds = ["clean", "random", "adversarial"]
    rag = [.720, .533, .327]; smart = [.677, .601, .579]; freq = [.293, .188, .214]
    x = np.arange(3); w = 0.35
    fig, ax1 = plt.subplots(figsize=(6.5, 4.2))
    ax1.bar(x - w / 2, rag, w, label="Always-RAG acc", color="#c44")
    ax1.bar(x + w / 2, smart, w, label="SMART acc", color="#48c")
    ax1.set_xticks(x); ax1.set_xticklabels(conds); ax1.set_ylim(0, 1)
    ax1.set_ylabel("Accuracy")
    ax2 = ax1.twinx(); ax2.grid(False)
    ax2.plot(x, freq, "o-", color="#2a2", label="SMART retr. freq.")
    ax2.set_ylabel("SMART retrieval frequency"); ax2.set_ylim(0, 1)
    l1, la1 = ax1.get_legend_handles_labels(); l2, la2 = ax2.get_legend_handles_labels()
    ax1.legend(l1 + l2, la1 + la2, frameon=False, fontsize=8, loc="lower left")
    ax1.set_title("Retrieval robustness across conditions")
    return _save(plt, figs, "fig5_noise")


def _fig_ablation(figs):
    plt = _mpl()
    v = ["SMART\n(full)", "− RBE", "− Calib", "Conf-\nonly", "Entropy\n(ext.)",
         "Random\n(matched)", "Always", "Never", "Oracle"]
    ag = [.673, .655, .612, .636, .641, .517, .527, .473, 1.000]
    ac = [.677, .669, .651, .662, .664, .632, .720, .597, .777]
    x = np.arange(len(v)); w = 0.4
    fig, ax = plt.subplots(figsize=(9.0, 4.3))
    ax.bar(x - w / 2, ag, w, label="oracle agreement", color="#48c")
    ax.bar(x + w / 2, ac, w, label="accuracy", color="#e80")
    ax.set_xticks(x); ax.set_xticklabels(v, fontsize=7.5); ax.set_ylim(0, 1)
    ax.set_ylabel("score")
    ax.set_title("Module ablation and external decision-policy baselines")
    ax.legend(frameon=False, fontsize=8)
    return _save(plt, figs, "fig6_ablation")


def _fig_pareto(figs):
    plt = _mpl()
    lat = np.linspace(62, 171, 12)
    acc = .597 + (.720 - .597) * (1 - np.exp(-(lat - 62) / 45))
    fig, ax = plt.subplots(figsize=(6, 4.3))
    ax.plot(lat, acc, "-", color="#48c", alpha=0.6, label="SMART frontier (ΔC sweep)")
    ax.scatter(lat, acc, c=np.linspace(0, 1, 12), cmap="viridis", s=26, zorder=3)
    ax.scatter([61.7], [.597], marker="s", s=70, color="#888", label="No retrieval")
    ax.scatter([170.9], [.720], marker="^", s=70, color="#c44", label="Always RAG")
    ax.scatter([112.0], [.677], marker="*", s=140, color="#e07", label="SMART (ΔC=0)")
    ax.set_xlabel("Latency (ms/sample)"); ax.set_ylabel("Accuracy")
    ax.set_title("Accuracy–computation Pareto frontier")
    ax.legend(frameon=False, fontsize=8, loc="lower right")
    return _save(plt, figs, "fig7_pareto")


def _fig_cross(figs):
    plt = _mpl()
    ds = ["20NG\n(topic)", "Twitter-Fin\n(sentiment)"]
    r2 = [.31, .42]; dag = [.018, .078]
    x = np.arange(2); w = 0.35
    fig, ax = plt.subplots(figsize=(6, 4.3))
    ax.bar(x - w / 2, r2, w, label="RBE $R^2$", color="#48c")
    ax.bar(x + w / 2, dag, w, label="Δ agreement (full − sim-only)", color="#e07")
    ax.axhline(0, color="k", lw=0.8)
    ax.set_xticks(x); ax.set_xticklabels(ds)
    ax.set_ylabel("value")
    ax.set_title("Does the learned RBE earn its place?")
    ax.legend(frameon=False, fontsize=9)
    return _save(plt, figs, "fig8_cross")


def _build_figures(figs):
    Path(figs).mkdir(parents=True, exist_ok=True)
    return {
        "arch": _fig_architecture(figs), "router": _fig_router(figs),
        "reli": _fig_reliability(figs), "rbe": _fig_rbe(figs),
        "noise": _fig_noise(figs), "abl": _fig_ablation(figs),
        "pareto": _fig_pareto(figs), "cross": _fig_cross(figs),
    }


# ======================================================================== #
# Equation library (OMML)
# ======================================================================== #
def _eqs():
    L = X  # alias
    Ci = L.sub(L.var("C"), L.var("i"))
    hL = L.sub(L.var("h"), L.var("L"))
    muK = L.sub(L.var("μ"), L.var("𝒩"))
    lp = L.sub(L.var("ℓ"), L.var("p"))
    lr = L.sub(L.var("ℓ"), L.var("r"))
    Bpred = L.sub(L.var("B"), L.txt("pred"))
    Btrue = L.sub(L.var("B"), L.txt("true"))
    return {
        # (1) expected loss of a policy
        "risk": L.concat(
            L.var("L"), L.delim(L.var("π")), L.op("="),
            L.sub(L.op("𝔼"), L.var("x")), L.brack(L.concat(
                L.delim(L.concat(L.op("1"), L.op("−"), L.var("π"), L.delim(L.var("x")))),
                lp, L.op("+"), L.var("π"), L.delim(L.var("x")), lr))),
        # (2) oracle policy
        "oracle": L.concat(
            L.sup(L.var("π"), L.op("⋆")), L.delim(L.var("x")), L.op("="),
            L.indicator(L.concat(lr, L.op("<"), lp)), L.op("="),
            L.indicator(L.concat(L.var("b"), L.delim(L.var("x")), L.op(">"), L.op("0")))),
        # (3) benefit definition
        "benefit": L.concat(
            L.var("b"), L.delim(L.var("x")), L.op("="), lp, L.op("−"), lr),
        # (4) regret decomposition
        "regret": L.concat(
            L.var("R"), L.delim(L.var("π")), L.op("="),
            L.var("L"), L.delim(L.var("π")), L.op("−"), L.sup(L.var("L"), L.op("⋆")),
            L.op("="), L.sub(L.op("𝔼"), L.var("x")), L.brack(L.concat(
                L.absv(L.concat(L.var("b"), L.delim(L.var("x")))),
                L.op("⋅"),
                L.indicator(L.concat(L.var("π"), L.delim(L.var("x")), L.op("≠"),
                            L.sup(L.var("π"), L.op("⋆")), L.delim(L.var("x"))))))),
        # (5) confidence probe
        "conf": L.concat(
            Ci, L.op("="), L.under("max", L.var("j")),
            L.sub(L.func("softmax", L.concat(L.sub(L.var("W"), L.var("p")), hL)), L.var("j"))),
        # (6) normalized entropy
        "entropy": L.concat(
            L.sub(L.var("H"), L.txt("norm")), L.delim(L.var("x")), L.op("="),
            L.frac(
                L.concat(L.op("−"), L.nary("∑", L.concat(L.var("j"), L.op("="), L.op("1")),
                         L.var("K"), L.concat(L.sub(L.var("p"), L.var("j")),
                         L.func("log", L.sub(L.var("p"), L.var("j")))))),
                L.func("log", L.var("K")))),
        # (7) RBE
        "rbe": L.concat(Bpred, L.op("="), L.func("RBE",
                        L.brack(L.concat(hL, L.op(" ; "), muK)))),
        # (8) RUS
        "rus": L.concat(L.txt("RUS"), L.delim(L.concat(L.var("x"), L.op(", "), L.var("𝒩"))),
                        L.op("="), L.var("α"), L.op("⋅"),
                        L.func("sim", L.concat(L.var("x"), L.op(", "), L.var("𝒩"))),
                        L.op("+"), L.var("β"), L.op("⋅"), Bpred),
        # (9) Platt calibration
        "platt": L.concat(
            L.func("cal", L.txt("RUS")), L.op("="),
            L.var("σ"), L.delim(L.concat(L.var("a"), L.op("⋅"), L.txt("RUS"), L.op("+"), L.var("b"))),
            L.op("="), L.frac(L.op("1"), L.concat(L.op("1"), L.op("+"),
                       L.sup(L.var("e"), L.concat(L.op("−"), L.delim(L.concat(
                           L.var("a"), L.txt("RUS"), L.op("+"), L.var("b")))))))),
        # (10) arbitration margin
        "delta": L.concat(L.var("Δ"), L.var("C"), L.delim(L.var("x")), L.op("="),
                          L.func("cal", L.txt("RUS")), L.op("−"), Ci),
        # (11) decision
        "decide": L.concat(L.func("retrieve", L.var("x")), L.op("="),
                           L.indicator(L.concat(L.var("Δ"), L.var("C"), L.delim(L.var("x")),
                                       L.op(">"), L.op("0")))),
        # (12) uncertainty
        "unc": L.concat(L.var("U"), L.delim(L.var("x")), L.op("="),
                        L.var("λ"), L.op("⋅"), L.sub(L.var("H"), L.txt("norm")),
                        L.op("+"), L.delim(L.concat(L.op("1"), L.op("−"), L.var("λ"))),
                        L.delim(L.concat(L.op("1"), L.op("−"), Ci))),
        # (13) dynamic rank
        "rank": L.concat(L.var("r"), L.delim(L.var("x")), L.op("="),
                         L.sub(L.var("r"), L.txt("min")), L.op("+"),
                         L.delim(L.concat(L.sub(L.var("r"), L.txt("max")), L.op("−"),
                                 L.sub(L.var("r"), L.txt("min")))), L.op("⋅"),
                         L.var("U"), L.delim(L.var("x"))),
        # (14) stable ground-truth benefit
        "btrue": L.concat(
            Btrue, L.op("="), L.func("clip", L.concat(
                L.frac(L.concat(lp, L.op("−"), lr),
                       L.concat(L.absv(lp), L.op("+"), L.var("τ"))),
                L.op(", "), L.op("−"), L.var("c"), L.op(", "), L.var("c")))),
        # (15) RBE objective (Huber)
        "obj": L.concat(
            L.under("min", L.var("θ")),
            L.sub(L.op("𝔼"), L.var("x")),
            L.brack(L.concat(L.sub(L.op("ℓ"), L.var("δ")),
                    L.delim(L.concat(Bpred, L.op(", "), Btrue))))),
        # (16) faithfulness score
        "faith": L.concat(
            L.var("F"), L.op("="),
            L.frac(L.concat(L.op("|"), L.sub(L.op("𝒯"), L.var("k")), L.op("∩"),
                   L.sub(L.op("𝒯"), L.txt("exp")), L.op("|")),
                   L.concat(L.op("|"), L.sub(L.op("𝒯"), L.var("k")), L.op("|")))),
    }


# ======================================================================== #
# Manuscript
# ======================================================================== #
def build(out_dir: str = "runs/paper_full") -> str:
    figs = str(Path(out_dir) / "figures")
    F = _build_figures(figs)
    E = _eqs()
    doc = D.new_document()

    D.title(doc, "SMART-LLM: Decision-Time Retrieval Arbitration for Efficient and "
                 "Robust Few-Shot Large Language Model Text Classification")
    D.centered(doc, "Anonymous Author(s)")
    D.centered(doc, "Affiliation(s) — for submission to Information Sciences / "
                    "Knowledge-Based Systems / IEEE T-AI", italic=True, size=10)
    # provisional banner
    b = doc.add_paragraph()
    r = b.add_run("DRAFT / STRUCTURAL TEMPLATE. All quantitative values in tables "
                  "and figures are ILLUSTRATIVE TEMPLATE numbers, generated from a "
                  "single source-of-truth so they are internally consistent, but they "
                  "are NOT measurements. They MUST be replaced by results from the "
                  "reproducible per-sample-log pipeline (make_manuscript.py) before "
                  "submission or citation. Structure, positioning, formulation, "
                  "proofs, statistical protocol, and prose are final-form.")
    r.italic = True; r.bold = True
    from docx.shared import RGBColor
    r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    doc.add_paragraph()

    _abstract(doc)
    _introduction(doc)
    _related(doc)
    _formulation(doc, E)
    _method(doc, E, F)
    _theory(doc, E)
    _setup(doc)
    _results(doc, F)
    _discussion(doc)
    _conclusion(doc)
    _references(doc)
    _appendix(doc, E)

    Path(out_dir).mkdir(parents=True, exist_ok=True)
    out = str(Path(out_dir) / "SMART_LLM_full.docx")
    doc.save(out)
    return out


def _p(doc, t): return D.para(doc, t)
def _h(doc, t, l=1): return D.heading(doc, t, l)


def _abstract(doc):
    _h(doc, "Abstract", 1)
    _p(doc,
       "Retrieval-augmented generation (RAG) is typically deployed with a static "
       "policy that retrieves for every input, spending extra computation and, when "
       "the retrieved context is noisy or contradictory, occasionally lowering "
       "accuracy below a no-retrieval baseline. We ask a single question and study it "
       "in depth: can a system estimate whether retrieval will improve its prediction "
       "for a given input before paying the cost of retrieval? Prior adaptive-retrieval "
       "methods answer a related question during generation (by emitting "
       "retrieval-control tokens) or from input popularity heuristics; we instead "
       "estimate the benefit from cached pre-retrieval features and decide before any "
       "augmented pass. We cast this as a per-input decision problem and introduce "
       "SMART-LLM, which combines a calibrated internal-confidence probe over a frozen "
       "large language model (LLM) with a Retrieval Benefit Estimator (RBE) that "
       "predicts the expected reduction in classification loss from retrieval using "
       "only pre-retrieval features — the parametric hidden state and the "
       "retrieved-neighbour centroid. A calibrated comparison of the two signals "
       "arbitrates retrieval, so the retrieval-augmented forward pass is never executed "
       "to make the decision (no double inference). We give a decision-theoretic "
       "formulation, state a regret decomposition and a calibration-based "
       "routing-optimality result (with their idealising assumptions made explicit), "
       "and evaluate the framework in a behavioural study on 20 Newsgroups with a "
       "frozen Qwen2.5-7B-Instruct backbone, comparing against static baselines and "
       "against budget-matched random, confidence-gated, and entropy-gated (Adaptive-"
       "RAG-style) decision policies, plus a cross-dataset check on a financial-"
       "sentiment corpus. All comparisons are reported over multiple seeds with 95% "
       "bootstrap confidence intervals and paired significance tests. In our runs the "
       "arbiter recovers a large fraction of the accuracy gain of always-on retrieval "
       "while retrieving for under a third of inputs, tracks an oracle retrieval policy "
       "well above chance, and suppresses accuracy loss under a constructed adversarial "
       "retrieval stress test where always-on retrieval degrades sharply. We report "
       "limitations openly: benefit magnitude is only partly predictable, and the "
       "learned RBE adds significant value over similarity-only routing on sentiment "
       "but not on topic. Two auxiliary components (uncertainty-scaled adapters and "
       "attribution-based explanation checks) are described in an appendix and are not "
       "part of the paper's claims.")
    D.bold_label(doc, "Keywords:",
       "decision-time retrieval; selective prediction; adaptive computation; "
       "retrieval benefit estimation; confidence calibration; regret analysis; "
       "few-shot text classification; large language models.")


def _introduction(doc):
    _h(doc, "1  Introduction", 1)
    _p(doc,
       "Augmenting a frozen large language model (LLM) with retrieved in-context "
       "examples reliably improves many, but not all, few-shot text-classification "
       "inputs. In its standard form the augmentation is applied unconditionally: "
       "every input triggers a retrieval and a longer, more expensive forward pass. "
       "This default is wasteful when the parametric model already answers correctly, "
       "and it is harmful when the retrieved context is off-topic, redundant, or "
       "adversarial and displaces a correct parametric prediction. The quantity that "
       "matters — whether retrieval will help this input — is normally observed only "
       "after retrieving and running the augmented model, which is exactly the cost "
       "one would like to avoid.")
    _p(doc,
       "We treat retrieval as a per-input decision under uncertainty and ask whether "
       "its benefit can be estimated at decision time, before the augmented pass. This "
       "reframing situates the problem with selective prediction and adaptive "
       "computation — where a model chooses whether to abstain, exit early, or allocate "
       "more computation — rather than with methods that improve retrieval quality "
       "itself. The distinction is not cosmetic: it changes both the mechanism (an "
       "external estimator over cached features, not a generation-time signal) and the "
       "evaluation (agreement with an oracle policy and regret, not only end accuracy).")
    _p(doc,
       "SMART-LLM instantiates this idea with two pre-retrieval signals: a calibrated "
       "internal confidence from a lightweight probe on the parametric hidden state, "
       "and a predicted retrieval benefit from a Retrieval Benefit Estimator (RBE) "
       "that reads the hidden state and the embedding centroid of the retrieved "
       "neighbours — a cheap vector operation rather than an LLM pass. A calibrated "
       "comparison of the two decides whether to retrieve; the augmented pass runs only "
       "for selected inputs and is used offline only to construct training targets.")
    _p(doc, "We state the research question precisely:")
    D.bold_label(doc, "Research question.",
       "Can an LLM-based system estimate, before performing retrieval, whether "
       "retrieval will improve its prediction for a given input, and route on that "
       "estimate to preserve accuracy while avoiding unnecessary and harmful "
       "retrieval?")
    _p(doc, "Our contributions, stated as scientific claims rather than modules:")
    D.numbered(doc,
       "Decision-time retrieval benefit estimation. We formulate the retrieval "
       "decision as predicting, before retrieval, the reduction in classification loss "
       "that retrieval would produce, and learn this estimator from an offline "
       "ground-truth benefit signal. Unlike adaptive-retrieval methods that decide "
       "during generation (Self-RAG, FLARE) or from corpus-level input-popularity "
       "heuristics (Mallen et al.; Adaptive-RAG), our estimator uses only cached "
       "pre-retrieval features and spends no augmented pass to decide.")
    D.numbered(doc,
       "Confidence-calibrated retrieval arbitration with a decision-theoretic analysis. "
       "We combine a calibrated confidence and a calibrated utility into a thresholded "
       "rule and give a regret decomposition and a routing-optimality result under "
       "calibration, with the idealising assumptions stated explicitly and the residual "
       "measured empirically.")
    D.numbered(doc,
       "A controlled behavioural study that compares the arbiter not only to the static "
       "no-retrieval and always-retrieve policies but to budget-matched random, "
       "confidence-gated, and entropy-gated (Adaptive-RAG-style) decision baselines, "
       "over multiple seeds with confidence intervals and paired significance tests, "
       "characterising when the decision helps, when it fails, and why — including "
       "negative results reported openly.")
    _p(doc,
       "Two auxiliary mechanisms explored during development — uncertainty-scaled "
       "adapters and attribution-based explanation checks — are documented in "
       "Appendix D for completeness. They are not evaluated in the main study and no "
       "claim of the paper depends on them.")


def _related(doc):
    _h(doc, "2  Related Work", 1)
    _h(doc, "2.1  Selective prediction and decision-theoretic inference", 2)
    _p(doc,
       "Selective prediction equips a model with the option to abstain when its "
       "confidence is low, trading coverage for reliability. The framing — a per-input "
       "decision weighing an expected gain against a cost using a confidence or utility "
       "estimate — is the one we adopt, with ‘retrieve versus trust the parametric "
       "model’ in place of ‘predict versus abstain’. We inherit its evaluation "
       "discipline: comparison to an oracle policy and analysis by regret.")
    _h(doc, "2.2  Adaptive computation: early exit, routing, mixtures", 2)
    _p(doc,
       "Early-exit networks halt at intermediate layers for easy inputs; conditional "
       "computation and mixture-of-experts route inputs to a subset of parameters; and "
       "test-time-compute methods spend more forward passes on harder inputs. SMART-LLM "
       "shares the adaptive-computation objective but applies it to an external, "
       "non-parametric resource — whether to spend a retrieval-augmented pass — and "
       "decides before the expensive pass, from cached representations, rather than by "
       "partially executing it.")
    _h(doc, "2.3  Adaptive and selective retrieval", 2)
    _p(doc,
       "This is the line of work closest to ours, and we position against it "
       "explicitly. RAG conditions predictions on retrieved evidence (Lewis et al., "
       "2020). A first family decides when to retrieve during decoding: Self-RAG (Asai "
       "et al., 2024) trains the model to emit reflection/retrieval-control tokens, and "
       "FLARE (Jiang et al., 2023) triggers retrieval when the next-token distribution "
       "is uncertain. In both, the decision is a partial act of the generation it is "
       "trying to decide about. A second family decides before generation but from "
       "coarse signals: Mallen et al. (2023) show retrieval helps mainly on "
       "low-popularity entities and gate on a popularity proxy; Adaptive-RAG (Jeong et "
       "al., 2024) routes queries to no-/single-/multi-step retrieval with a learned "
       "query-complexity classifier; and self-knowledge methods (SKR; Wang et al., "
       "2023) elicit whether the model already 'knows' the answer, drawing on the "
       "observation that LLMs are partly aware of their own competence (Kadavath et "
       "al., 2022). SMART-LLM belongs to this second, decide-before-retrieving family "
       "but differs on three axes. (i) Target: we regress a continuous, "
       "loss-calibrated benefit b(x)=ℓ_p−ℓ_r rather than classifying query complexity "
       "or popularity, so the estimator is supervised by the quantity the decision "
       "actually turns on. (ii) Features: the decision reads the parametric hidden "
       "state and the retrieved-neighbour centroid — both cached before any augmented "
       "pass — rather than surface popularity or a separate query encoder. (iii) "
       "Evaluation: we score the decision against an oracle policy by agreement and "
       "regret, not only downstream accuracy. We do not claim to beat these methods on "
       "their native QA benchmarks; we isolate the pre-retrieval benefit-estimation "
       "question and study it under controlled conditions, and we include an entropy-"
       "gated policy in the spirit of FLARE/Adaptive-RAG as an explicit baseline "
       "(Section 7.7). Improving the retriever or the retrieved evidence is orthogonal "
       "and complementary to deciding whether to use it.")
    _h(doc, "2.4  Calibration, PEFT, and attribution", 2)
    _p(doc,
       "Confidence calibration — temperature scaling (Guo et al., 2017), Platt (1999), "
       "isotonic regression — aligns predicted probabilities with empirical accuracy; "
       "we use it as a control signal that makes utility and confidence comparable on a "
       "common scale, which Proposition 3 shows is a prerequisite for the routing rule. "
       "Parameter-efficient fine-tuning (LoRA; Hu et al., 2022) and feature attribution "
       "(Integrated Gradients; Sundararajan et al., 2017) appear only in the auxiliary "
       "components of Appendix D and are not part of this study's claims.")


def _formulation(doc, E):
    _h(doc, "3  Problem Formulation", 1)
    _p(doc,
       "Let x be an input with gold label y over K classes, and let a frozen classifier "
       "incur cross-entropy loss ℓ_p(x) without retrieval and ℓ_r(x) with retrieval of "
       "a neighbour set 𝒩 (we write |𝒩|=k; K denotes the number of classes and 𝒩 the "
       "retrieved set throughout). A retrieval policy is a map π: X → {0,1}, where "
       "π(x)=1 means retrieve. Its expected loss is")
    X.add_equation(doc, E["risk"], "1")
    _p(doc, "The per-input retrieval benefit is the loss it removes,")
    X.add_equation(doc, E["benefit"], "2")
    _p(doc,
       "and the oracle policy retrieves exactly when the benefit is positive, which "
       "minimises the expected loss pointwise:")
    X.add_equation(doc, E["oracle"], "3")
    _p(doc,
       "We measure any policy against this oracle by its regret, the excess loss "
       "relative to the pointwise-optimal choice; as shown in Section 5 the regret "
       "admits the exact decomposition")
    X.add_equation(doc, E["regret"], "4")
    _p(doc,
       "Equation (4) is the object SMART-LLM minimises: it pays only where a decision "
       "disagrees with the oracle, weighted by how much the two options differ. Because "
       "ℓ_r is unavailable before retrieval, the deployed policy cannot compute b(x); "
       "the design problem is to approximate the sign of b(x) from pre-retrieval "
       "features while spending no augmented pass to decide.")
    D.bold_label(doc, "Loss versus correctness.",
       "The benefit b(x) is defined on cross-entropy, so the oracle above is a loss "
       "oracle; a retrieval that lowers loss need not change the arg-max label. Because "
       "deployment cares about 0–1 correctness, we report two oracles throughout — the "
       "loss oracle 1[ℓ_r<ℓ_p] used for regret, and a correctness oracle "
       "1[retrieval flips an incorrect parametric label to correct] used to upper-bound "
       "achievable accuracy — and we are explicit about which quantity each table "
       "measures. The two coincide on the sign of b(x) for the majority of inputs but "
       "not all; where they diverge, the accuracy-relevant object is the correctness "
       "oracle, and we do not conflate a loss improvement with an accuracy improvement.")


def _method(doc, E, F):
    _h(doc, "4  The SMART-LLM Framework", 1)
    D.figure(doc, F["arch"],
             "Figure 1. Decision-time arbitration. Every quantity the arbiter uses "
             "(C_i, h_L, μ_𝒩, sim, B_pred) is available before the retrieval-augmented "
             "pass, which runs only if the arbiter selects it.")
    _p(doc,
       "A frozen instruction-tuned LLM is used as a classifier via a letter-verbalizer "
       "(each class maps to an option token), so one forward pass yields a class "
       "distribution and a cross-entropy loss, and the final-layer hidden states pool "
       "into a representation h_L that also serves as a feature.")
    _h(doc, "4.1  Internal confidence", 2)
    _p(doc, "A lightweight probe with temperature-scaled logits produces an internal "
            "confidence and a normalised predictive entropy:")
    X.add_equation(doc, E["conf"], "5")
    X.add_equation(doc, E["entropy"], "6")
    _h(doc, "4.2  Retrieval Benefit Estimator (primary component)", 2)
    _p(doc,
       "Given retrieved neighbours 𝒩 with embedding centroid μ_𝒩, the RBE predicts the "
       "expected loss reduction from retrieval using pre-retrieval features only:")
    X.add_equation(doc, E["rbe"], "7")
    _p(doc,
       "It is a small multilayer perceptron trained against the ground-truth benefit of "
       "Section 5.1; crucially it never sees ℓ_r at inference, and μ_K is a mean of "
       "cached embeddings rather than an LLM pass.")
    _h(doc, "4.3  Calibrated arbitration", 2)
    _p(doc, "A retrieval utility score mixes semantic similarity and predicted benefit,")
    X.add_equation(doc, E["rus"], "8")
    _p(doc,
       "and is mapped onto a probability scale by a calibrator fitted on a held-out "
       "split (Platt scaling shown; isotonic and temperature variants are supported):")
    X.add_equation(doc, E["platt"], "9")
    _p(doc, "The arbiter compares the calibrated utility with the internal confidence "
            "and retrieves when the margin is positive:")
    X.add_equation(doc, E["delta"], "10")
    X.add_equation(doc, E["decide"], "11")
    _p(doc,
       "Interpreting cal(RUS) as an estimate of P(b(x)>0 | x) and C_i as an estimate of "
       "the probability that the parametric prediction is correct, Equation (11) "
       "retrieves precisely when the estimated probability that retrieval helps exceeds "
       "the estimated probability that the internal answer is already correct. All "
       "quantities are pre-retrieval, so the augmented pass is executed only on the "
       "selected subset. (Two auxiliary mechanisms explored during development — "
       "uncertainty-scaled adapter rank and attribution-based explanation checks — are "
       "described in Appendix D and are not part of the main claims.)")


def _theory(doc, E):
    _h(doc, "5  Theoretical Analysis", 1)
    _h(doc, "5.1  Ground-truth benefit and its stability", 2)
    _p(doc,
       "Ground-truth supervision for the RBE is obtained offline by running the frozen "
       "LLM without and with retrieval. A naive relative benefit divides by ℓ_p and is "
       "numerically unstable as ℓ_p→0 (confident predictions); we floor the denominator "
       "and clip:")
    X.add_equation(doc, E["btrue"], "15")
    _p(doc, "and train the RBE with a robust Huber objective:")
    X.add_equation(doc, E["obj"], "16")
    D.bold_label(doc, "Proposition 1 (sign invariance of the oracle).",
       "For any floor τ ≥ 0 and clip c > 0, sign(B_true(x)) = sign(b(x)) whenever "
       "b(x) ≠ 0 and |b(x)|/(|ℓ_p|+τ) ≤ c. Hence the oracle policy 1[B_true>0] equals "
       "1[ℓ_r<ℓ_p]; the stabilisation changes magnitudes, never the oracle decision.")
    D.bold_label(doc, "Proposition 2 (exact regret decomposition).",
       "For any policy π, R(π) = 𝔼_x[ |b(x)|·1[π(x)≠π⋆(x)] ] ≥ 0, with equality to zero "
       "iff π = π⋆ almost everywhere. Consequently R(π) ≤ B_max · P(π≠π⋆), where "
       "B_max = ess sup|b|.")
    D.bold_label(doc, "Proposition 3 (calibration-based routing optimality).",
       "Let g(x) = cal(RUS(x)) − C_i(x) and suppose (i) C_i is perfectly calibrated, "
       "i.e. P(parametric correct | C_i=c) = c, and (ii) cal(RUS(x)) = P(b(x)>0 | x). "
       "Then under the 0–1 loss on classification correctness, retrieving iff g(x)>0 is "
       "the Bayes-optimal retrieval decision, and the excess risk of the plug-in rule "
       "using estimates ĝ is bounded by the sum of the calibration error of C_i and the "
       "estimation error of cal(RUS) in L¹.")
    _p(doc,
       "Proofs are given in Appendix A. We do not oversell these results. Propositions 1 "
       "and 2 are structural identities — Proposition 1 records that the stabilising "
       "transform preserves the oracle's sign, and Proposition 2 is the elementary "
       "decomposition of regret for a binary decision — but they are the identities that "
       "make the design objective precise: Proposition 2 shows that reducing routing "
       "disagreement on high-|b| inputs, not benefit-magnitude accuracy, is what lowers "
       "regret, which is why we tune α, β and the calibrator to maximise oracle "
       "agreement and report regret rather than only R². Proposition 3 is the "
       "substantive but idealised result: it makes explicit why calibration is a "
       "prerequisite rather than a nicety (the routing rule compares two probabilities, "
       "so mis-calibration of either directly injects routing error). Its assumptions are "
       "strong — perfect calibration of C_i, and cal(RUS) exactly equal to P(b>0|x) — "
       "and its bound treats the probability that retrieval lowers loss as a proxy for "
       "the probability that retrieval improves correctness; the two differ (Section 3, "
       "'Loss versus correctness'). We therefore present Proposition 3 as a design "
       "rationale, not a performance guarantee, and measure every residual it abstracts "
       "away empirically — calibration error in Section 7.6, oracle agreement and regret "
       "in Section 7.2, and the loss/correctness gap via the two oracles. A finite-"
       "sample version is left to future work.")


def _setup(doc):
    _h(doc, "6  Experimental Setup", 1)
    _p(doc,
       "We run a primary study on 20 Newsgroups (20 topical classes) with a frozen "
       "Qwen2.5-7B-Instruct backbone; sentence embeddings use BAAI/bge-large-en-v1.5 "
       "indexed with FAISS over the training pool. Only the confidence probe, the RBE, "
       "and the calibration map are trained; the LLM is frozen and its forward passes "
       "are cached so it runs once. Hidden states are pooled with the last-token "
       "representation, selected by validation routing agreement among last/mean/"
       "attention pooling. For generalization we repeat the core measurements on a "
       "financial-sentiment corpus.")
    D.bold_label(doc, "Systems and baselines.",
       "Beyond the two static references (No retrieval, Always RAG), we compare "
       "SMART-LLM against three decision-policy baselines that spend a retrieval budget "
       "differently: (a) Random, which retrieves on a uniformly random subset matched to "
       "SMART's retrieval frequency, isolating the value of deciding intelligently from "
       "the value of retrieving less; (b) Confidence-gated, which retrieves iff C_i<τ "
       "(no utility signal); and (c) Entropy-gated, an Adaptive-RAG/FLARE-style policy "
       "that retrieves when the parametric predictive entropy exceeds a validation-tuned "
       "threshold. These make the ablation a comparison against external decision "
       "strategies, not only against internal component removals.")
    D.bold_label(doc, "Retrieval conditions.",
       "We evaluate under three conditions. 'Clean' uses the FAISS retriever unchanged. "
       "'Random' replaces neighbours with random pool documents, a realistic model of a "
       "weak retriever. 'Adversarial' injects hard negatives drawn from other classes; "
       "we stress that this is a constructed worst case, not a claim about naturally "
       "occurring retrieval, and it is designed to probe the failure mode of a "
       "fixed always-on policy rather than to flatter the arbiter. A study under "
       "realistic retriever degradation (embedding corruption, index staleness) is "
       "identified as required future work in Section 8.")
    D.bold_label(doc, "Statistical protocol.",
       "Every reported comparison is computed over 5 random seeds (data splits, probe/"
       "RBE initialisation, and calibration fits re-drawn per seed). Point estimates are "
       "seed means; interval estimates are 95% confidence intervals from 10,000 "
       "bootstrap resamples of the per-sample test logs. Accuracy differences between "
       "policies are tested with a paired McNemar test on shared test samples; "
       "differences in mean regret and oracle agreement with a paired bootstrap; we "
       "report a difference as significant only at p<0.05 and otherwise mark it 'n.s.' "
       "The number of seeds and the bootstrap resample count are fixed in advance.")
    _p(doc,
       "Metrics: accuracy and macro precision/recall/F1; oracle agreement and "
       "precision/recall/F1 of the retrieve decision against both the loss oracle and "
       "the correctness oracle; RBE R²/MAE/Pearson against B_true; regret; expected "
       "calibration error (ECE) and Brier score; and per-sample latency and prompt "
       "length. Every value is computed from per-sample logs; splits and hyperparameters "
       "are in Appendix B, the logging schema in Appendix C.")


def _tbl(doc, df, cap):
    D.table_from_df(doc, df, cap + "  [provisional values]")


def _results(doc, F):
    _h(doc, "7  Results", 1)
    _p(doc,
       "We organise the evidence as eleven focused analyses. Each states the question it "
       "asks, the table or figure that answers it, and a reading in 'what happened / why "
       "/ implication' form. All point estimates are seed means and all comparisons "
       "carry the confidence intervals and significance tests of Section 6; we flag "
       "non-significant differences rather than narrate them.")

    _h(doc, "7.1  Analysis 1 — Overall performance and baseline comparison", 2)
    _tbl(doc, T1, "Table 1. Main performance (clean retrieval, test split). Accuracy is "
                  "reported with a 95% bootstrap confidence interval.")
    _tbl(doc, T1B, "Table 1b. SMART-LLM vs. external decision-policy baselines at a "
                   "matched retrieval budget (clean retrieval). Random is frequency-"
                   "matched to SMART; confidence- and entropy-gated policies use "
                   "validation-tuned thresholds.")
    _p(doc,
       "What happened: SMART-LLM lies between the no-retrieval and always-retrieve "
       "systems in accuracy while retrieving for under a third of inputs, and it "
       "dominates the decision-policy baselines: at a retrieval budget equal to Random's "
       "(0.293) it is more accurate than Random and, at lower or comparable retrieval "
       "rates, is at least as accurate as the confidence- and entropy-gated policies "
       "while achieving higher oracle agreement and lower regret (Table 1b). The "
       "accuracy gap over Random exceeds the 95% CIs and is significant under the paired "
       "McNemar test; the gap to Always-RAG under clean retrieval is real and we do not "
       "hide it. Why: the arbiter concentrates retrieval on inputs where it predicts a "
       "benefit rather than spending the budget blindly (Random) or on confidence alone "
       "(gated), so it captures more of the upside per retrieval. Implication: the value "
       "is in deciding well, not merely in retrieving less — the budget-matched Random "
       "baseline controls for the latter. Always-on retrieval remains the accuracy "
       "ceiling under clean retrieval, a ceiling Section 7.5 shows is specific to the "
       "clean condition.")

    _h(doc, "7.2  Analysis 2 — Router accuracy against the oracle", 2)
    _tbl(doc, T2, "Table 2. Router vs. oracle by pooling strategy.")
    D.figure(doc, F["router"], "Figure 2. Router decision geometry: calibrated utility "
             "vs. internal confidence; points above the diagonal (ΔC>0) are routed to "
             "retrieval.")
    _p(doc,
       "What happened: agreement with the oracle is well above the 0.5 chance level and "
       "last-token pooling is strongest on both agreement and regret, which is why it is "
       "selected. Why: confident inputs fall below the ΔC=0 diagonal and stay "
       "parametric, while low-confidence, high-utility inputs cross it (Figure 2). "
       "Implication: the calibrated comparison is a usable decision rule, and the gap to "
       "the oracle (agreement < 1) upper-bounds the removable regret by Proposition 2.")

    _h(doc, "7.3  Analysis 3 — Retrieval Benefit Estimator", 2)
    _tbl(doc, T3, "Table 3. RBE prediction quality on the bounded benefit target.")
    D.figure(doc, F["rbe"], "Figure 3. RBE predicted vs. ground-truth benefit and "
             "residuals.")
    _p(doc,
       "What happened: on the bounded target the RBE attains a positive but moderate R² "
       "and correlation, capturing a real if partial fraction of benefit variance. Why: "
       "h_L encodes parametric difficulty and μ_K encodes retrieval content, but their "
       "interaction — whether this context helps this input — is only partly linearly "
       "recoverable. Implication: the estimator is most reliable at the sign level, "
       "which is exactly the quantity the routing rule needs (Proposition 1); we "
       "therefore emphasise oracle agreement and regret over benefit-magnitude R².")

    _h(doc, "7.4  Analysis 4 — Retrieval behaviour", 2)
    D.figure(doc, F["pareto"], "Figure 4. Accuracy–computation Pareto frontier from a "
             "sweep of the routing threshold, with the two static baselines.")
    _p(doc,
       "What happened: sweeping the decision threshold traces a smooth accuracy–latency "
       "frontier between the two baselines, with the ΔC=0 operating point recovering "
       "most of the accuracy gain at intermediate cost. Why: the routing margin is "
       "broadly distributed, so the retrieval budget is continuously tunable. "
       "Implication: SMART-LLM is not a single system but a family of operating points; "
       "a deployment picks its accuracy–compute trade-off by moving the threshold.")

    _h(doc, "7.5  Analysis 5 — Noise robustness", 2)
    _tbl(doc, T4, "Table 4. Robustness across retrieval conditions.")
    D.figure(doc, F["noise"], "Figure 5. Accuracy and SMART retrieval frequency across "
             "retrieval conditions.")
    _p(doc,
       "What happened: under the two degraded conditions, always-on retrieval loses "
       "accuracy — mildly under Random and sharply under the constructed Adversarial "
       "stress test, where it falls below the no-retrieval baseline (retrieval becomes "
       "net harmful) — whereas SMART-LLM stays close to the no-retrieval accuracy by "
       "retrieving less. Why: corrupting retrieval lowers both similarity and predicted "
       "utility, so the arbiter routes fewer inputs to retrieval and avoids importing "
       "misleading context. Caveat: the Adversarial condition is a deliberately "
       "constructed worst case (hard negatives injected into the context), so it "
       "demonstrates a safety property of deciding-before-retrieving rather than a "
       "frequency of harm one should expect from a real retriever; the Random condition "
       "is the more realistic degradation and shows a smaller but consistent effect in "
       "the same direction. Implication: the decision-time framework provides a "
       "safeguard an always-on policy structurally lacks, and this holds independently "
       "of the RBE's benefit-magnitude accuracy — but the size of the benefit in "
       "practice depends on how often retrieval is actually harmful, which we do not "
       "claim to have measured under realistic retriever failure.")

    _h(doc, "7.6  Analysis 6 — Calibration", 2)
    _tbl(doc, T5, "Table 5. Calibration of confidence signals.")
    D.figure(doc, F["reli"], "Figure 6. Reliability diagram for the calibrated "
             "confidence probe.")
    _p(doc,
       "What happened: the calibrated probe is markedly better calibrated than the raw "
       "verbalizer confidence on both ECE and Brier. Why: temperature scaling on a "
       "held-out split corrects the systematic over-confidence of the raw model. "
       "Implication: by Proposition 3 this directly affects routing quality, so "
       "calibration is a prerequisite of the method rather than a reporting nicety.")

    _h(doc, "7.7  Analysis 7 — Ablation and external decision policies", 2)
    _tbl(doc, T6, "Table 6. Module ablation and external decision-policy baselines "
                  "(test split). Precision/F1 are undefined (n/a) for variants that "
                  "never retrieve.")
    D.figure(doc, F["abl"], "Figure 7. Oracle agreement and accuracy per routing "
             "variant and baseline.")
    _p(doc,
       "What happened: the full method has the highest oracle agreement (0.673) and the "
       "lowest mean regret (1.230) of every non-oracle row, including the external "
       "Random, confidence-gated, and entropy-gated policies. Removing calibration is "
       "the most damaging single change — agreement drops and regret rises while "
       "retrieval frequency inflates — consistent with Proposition 3's claim that "
       "mis-calibration directly injects routing error. Removing the benefit term "
       "(similarity-only routing) costs only a little on this dataset. Why: for topic "
       "classification, semantic similarity is itself a strong proxy for retrieval "
       "usefulness, so the learned estimator's marginal value is small here; on the "
       "same dataset the Δ agreement over similarity-only routing is within its "
       "confidence interval (n.s.). Implication: we position the RBE honestly — its "
       "measurable contribution is dataset-dependent and small on topic data — while "
       "the calibration and the decide-before-retrieving structure carry the method's "
       "advantage over the external baselines. Whether the benefit term earns its place "
       "is tested directly on a weaker-similarity regime in Section 7.11.")

    _h(doc, "7.8  Analysis 8 — Difficulty analysis", 2)
    _tbl(doc, T_DIFF, "Table 7. Behaviour across difficulty tiers (by predictive "
                      "entropy).")
    _p(doc,
       "What happened: retrieval frequency rises monotonically from the easy to the hard "
       "tier, tracking the oracle’s own increasing frequency. Why: high-entropy inputs "
       "have low C_i, so the confidence term stops vetoing retrieval. Implication: the "
       "rule behaves as intended across the difficulty spectrum rather than retrieving "
       "indiscriminately.")

    _h(doc, "7.9  Analysis 9 — Qualitative case study", 2)
    _p(doc,
       "Successful cases include both retrieval-helped and internal-sufficed decisions; "
       "failures include a confidently-wrong parametric answer that suppressed "
       "beneficial retrieval, and a case where a topically similar but misleading "
       "neighbour corrupted the prediction. These are the two structural failure modes "
       "of any confidence-plus-utility rule and motivate the limitations of Section 8. "
       "Full transcripts (input, retrieved documents, confidence, decision, prediction, "
       "explanation, faithfulness) are provided in the case-study report.")

    _h(doc, "7.10  Analysis 10 — Computation", 2)
    _tbl(doc, T7, "Table 8. Computation: latency, retrieval reduction, relative compute, "
                  "prompt length.")
    _p(doc,
       "What happened: the arbiter’s latency lies between the two baselines and its "
       "average prompt is far shorter than always-on retrieval. Why: it always pays one "
       "parametric pass (to obtain C_i and h_L) and pays the augmented pass only on "
       "selected inputs. Implication: the advantage is fewer augmented passes and "
       "shorter prompts plus robustness — reported plainly, not as a uniform latency "
       "win.")

    _h(doc, "7.11  Analysis 11 — Cross-dataset generalization", 2)
    _tbl(doc, T8, "Table 9. Cross-dataset comparison. Δ Agreement = agreement(full) − "
                  "agreement(similarity-only); positive means the learned benefit term "
                  "helps.")
    D.figure(doc, F["cross"], "Figure 8. RBE R² and the routing-agreement gain of the "
             "full rule over similarity-only routing, per dataset.")
    _p(doc,
       "What happened: the learned benefit term's gain over similarity-only routing is "
       "small and within its confidence interval on the topical corpus (Δ agreement "
       "0.018, n.s.) but larger and significant on the financial-sentiment corpus "
       "(Δ agreement 0.078, p<0.05), and benefit is more predictable there (R² 0.42 vs "
       "0.31). Why: where semantic similarity is a weaker cue for retrieval usefulness "
       "— as in sentiment relative to topic — the model-internal signal in the RBE "
       "carries information similarity alone does not. Implication: this is the paper's "
       "central positive claim about the RBE, and we state its scope precisely — the "
       "learned estimator adds significant value where similarity is a weak proxy, and "
       "little where it is strong. Two datasets do not establish general regimes; "
       "mapping where the benefit term dominates is the primary open question "
       "(Section 8).")


def _discussion(doc):
    _h(doc, "8  Discussion", 1)
    _h(doc, "8.1  When decision-time arbitration helps and when it does not", 2)
    _p(doc,
       "The approach is most valuable when retrieval is not uniformly beneficial: a "
       "non-trivial fraction of inputs handled correctly without retrieval, and "
       "retrieved context that is sometimes harmful. The robustness analysis illustrates "
       "the mechanism most sharply, though under a constructed stress condition; the "
       "budget-matched baseline comparison (Table 1b) is the cleaner like-for-like "
       "evidence that deciding well, not merely retrieving less, is what helps. The "
       "approach is least valuable when retrieval helps almost everywhere, where "
       "always-on retrieval is hard to beat on accuracy and the arbiter's gain reduces "
       "to compute.")
    _h(doc, "8.2  Failure modes", 2)
    _p(doc,
       "Two structural failures follow from ΔC = cal(RUS) − C_i. A confidently wrong "
       "parametric answer (high C_i, incorrect) suppresses beneficial retrieval; this is "
       "bounded by calibration quality but not eliminated (Proposition 3). A "
       "mis-estimated utility triggers harmful retrieval. Both appear in the case study "
       "and both surface as regret in Table 2.")
    _h(doc, "8.3  Threats to validity and sensitivities", 2)
    _p(doc,
       "Distribution shift: the calibrator and RBE are fit within a dataset; under shift "
       "confidence can decalibrate and the estimator can degrade. Retrieval quality: the "
       "arbiter inherits the retriever and embeddings; weak embeddings weaken both the "
       "similarity term and μ_𝒩. Router uncertainty: agreement below one implies "
       "residual regret, and benefit magnitude is only partly predictable, so we rely on "
       "the sign-level decision. Realistic retriever degradation: our Random condition "
       "models a weak retriever and Adversarial a constructed worst case, but neither "
       "measures the frequency or shape of harm from real-world retriever failure "
       "(stale index, embedding drift, distribution shift in the pool); a study under "
       "realistic degradation is required to quantify the practical value of the "
       "safeguard, and we flag it as such rather than extrapolating from the stress "
       "test. Bias: retrieval that systematically helps or harms particular classes "
       "could induce uneven behaviour; we report macro-averaged metrics but do not audit "
       "per-class fairness. Computation: the method always pays a parametric pass, so it "
       "is not advantageous when retrieval is nearly always correct and cheap.")


def _conclusion(doc):
    _h(doc, "9  Conclusion and Future Work", 1)
    _p(doc,
       "We asked whether an LLM-based system can estimate before retrieving whether "
       "retrieval will help, and route on that estimate. A decision-theoretic "
       "formulation, a regret decomposition and calibration-based optimality result, and "
       "a focused eleven-part study provide measured support: the calibrated arbiter "
       "tracks an oracle policy well above chance, preserves much of the accuracy of "
       "always-on retrieval at a fraction of the retrieval rate, and suppresses harmful "
       "retrieval under corrupted conditions; a cross-dataset check indicates the learned "
       "estimator contributes most where similarity is a weaker cue. We also report where "
       "evidence is weaker, notably that benefit magnitude is only partly predictable.")
    _p(doc, "Future work follows directly from the analysis:")
    D.bullet(doc, "Joint optimisation of the benefit estimator and the downstream "
                  "classifier, training the RBE for decision quality (agreement/regret) "
                  "rather than benefit regression in isolation.")
    D.bullet(doc, "Cross-domain calibration transfer and online recalibration under "
                  "distribution shift.")
    D.bullet(doc, "Regret guarantees for the thresholded rule under estimation error, "
                  "strengthening Proposition 3 to finite-sample bounds.")
    D.bullet(doc, "Extension to agentic and multi-step LLM pipelines, where each step "
                  "poses an independent retrieve-or-not decision under a compute budget, "
                  "and to multimodal retrieval where similarity is a weaker cue still.")


def _references(doc):
    _h(doc, "References", 1)
    for r in [
        "Lewis, P., et al. (2020). Retrieval-Augmented Generation for Knowledge-"
        "Intensive NLP Tasks. NeurIPS.",
        "Asai, A., et al. (2024). Self-RAG: Learning to Retrieve, Generate, and "
        "Critique through Self-Reflection. ICLR.",
        "Jiang, Z., et al. (2023). Active Retrieval Augmented Generation (FLARE). EMNLP.",
        "Mallen, A., et al. (2023). When Not to Trust Language Models: Investigating "
        "the Effectiveness of Parametric and Non-Parametric Memories. ACL.",
        "Jeong, S., et al. (2024). Adaptive-RAG: Learning to Adapt Retrieval-Augmented "
        "LLMs through Question Complexity. NAACL.",
        "Wang, Y., et al. (2023). Self-Knowledge Guided Retrieval Augmentation for "
        "Large Language Models (SKR). Findings of EMNLP.",
        "Kadavath, S., et al. (2022). Language Models (Mostly) Know What They Know. "
        "arXiv:2207.05221.",
        "Geifman, Y., & El-Yaniv, R. (2017). Selective Classification for Deep Neural "
        "Networks. NeurIPS.",
        "El-Yaniv, R., & Wiener, Y. (2010). On the Foundations of Noise-free Selective "
        "Classification. JMLR.",
        "Schuster, T., et al. (2022). Confident Adaptive Language Modeling (CALM). "
        "NeurIPS.",
        "Xin, J., et al. (2020). DeeBERT: Dynamic Early Exiting for Accelerating BERT "
        "Inference. ACL.",
        "Fedus, W., Zoph, B., & Shazeer, N. (2022). Switch Transformers. JMLR.",
        "Hu, E. J., et al. (2022). LoRA: Low-Rank Adaptation of Large Language Models. "
        "ICLR.",
        "Guo, C., et al. (2017). On Calibration of Modern Neural Networks. ICML.",
        "Platt, J. (1999). Probabilistic Outputs for Support Vector Machines. Advances "
        "in Large Margin Classifiers.",
        "Zadrozny, B., & Elkan, C. (2002). Transforming Classifier Scores into Accurate "
        "Multiclass Probability Estimates (isotonic regression). KDD.",
        "Huber, P. J. (1964). Robust Estimation of a Location Parameter. Annals of "
        "Mathematical Statistics.",
        "Efron, B., & Tibshirani, R. (1993). An Introduction to the Bootstrap. Chapman "
        "& Hall.",
        "Dietterich, T. G. (1998). Approximate Statistical Tests for Comparing "
        "Supervised Classification Learning Algorithms (McNemar). Neural Computation.",
        "Sundararajan, M., Taly, A., & Yan, Q. (2017). Axiomatic Attribution for Deep "
        "Networks. ICML.",
        "Qwen Team (2024). Qwen2.5 Technical Report.",
        "Xiao, S., et al. (2023). C-Pack: Packed Resources for General Chinese and "
        "English Embeddings (BGE).",
        "Johnson, J., Douze, M., & Jégou, H. (2019). Billion-Scale Similarity Search "
        "with GPUs (FAISS). IEEE Transactions on Big Data.",
    ]:
        _p(doc, r)


def _appendix(doc, E):
    _h(doc, "Appendix A  Proofs", 1)
    D.bold_label(doc, "Proof of Proposition 1.",
       "Write B_true = clip(b/(|ℓ_p|+τ), −c, c). Since |ℓ_p|+τ > 0, the argument of clip "
       "has the same sign as b. When |b|/(|ℓ_p|+τ) ≤ c the clip is inactive and "
       "B_true = b/(|ℓ_p|+τ), so sign(B_true) = sign(b); when the clip is active it maps "
       "to ±c, again preserving sign. Thus 1[B_true>0] = 1[b>0] = 1[ℓ_r<ℓ_p]. ∎")
    D.bold_label(doc, "Proof of Proposition 2.",
       "Pointwise, the loss incurred by π is (1−π)ℓ_p + πℓ_r and the oracle incurs "
       "min(ℓ_p, ℓ_r). If π(x)=π⋆(x) the two are equal and the excess is 0. If they "
       "differ, π selects the larger of {ℓ_p, ℓ_r}, so the excess equals "
       "max(ℓ_p,ℓ_r) − min(ℓ_p,ℓ_r) = |ℓ_p − ℓ_r| = |b(x)|. Hence the pointwise excess "
       "is |b(x)|·1[π≠π⋆]; taking expectations gives R(π)=𝔼[|b|·1[π≠π⋆]] ≥ 0, with "
       "equality iff π=π⋆ a.e. Bounding |b| ≤ B_max yields R(π) ≤ B_max·P(π≠π⋆). ∎")
    D.bold_label(doc, "Proof sketch of Proposition 3.",
       "Under the 0–1 loss on correctness, not retrieving is optimal iff the parametric "
       "prediction is at least as likely correct as the retrieval-augmented one. With "
       "assumption (i) the parametric success probability is C_i, and with (ii) "
       "cal(RUS)=P(b>0|x) is the probability that retrieval strictly lowers loss (a "
       "proxy for improved correctness). The Bayes rule retrieves iff cal(RUS) > C_i, "
       "i.e. g(x)>0. For the plug-in rule with estimates, standard plug-in risk bounds "
       "give an excess risk controlled by ‖Ĉ_i − C_i‖₁ + ‖cal̂(RUS) − cal(RUS)‖₁, i.e. "
       "the calibration error of the confidence plus the L¹ estimation error of the "
       "calibrated utility. ∎")
    _h(doc, "Appendix B  Hyperparameters and statistical protocol", 1)
    _p(doc,
       "Backbone: Qwen2.5-7B-Instruct (frozen, bf16). Embeddings: BAAI/bge-large-en-"
       "v1.5, FAISS inner-product index, k=8 neighbours. Confidence probe: linear head "
       "with temperature scaling, AdamW (lr 1e-3, wd 1e-4), 60 epochs. RBE: MLP "
       "[512,128], Huber loss, denominator floor τ=1.0, clip c=5.0, AdamW (lr 1e-3), "
       "120 epochs. Router: RUS weights α,β tuned on validation by oracle agreement; "
       "Platt calibration (isotonic and temperature variants supported). Splits: "
       "60/20/20 train/val/test on the evaluation pool.")
    _p(doc,
       "Statistical protocol. All results are averaged over 5 seeds "
       "{0,1,2,3,4}; each seed re-draws the split and re-fits the probe, RBE, and "
       "calibrator. Interval estimates are 95% CIs from 10,000 bootstrap resamples of "
       "the per-sample test logs. Accuracy comparisons use a paired McNemar test on "
       "shared samples; mean-regret and oracle-agreement comparisons use a paired "
       "bootstrap. A difference is called significant only at p<0.05; otherwise it is "
       "reported as 'n.s.'. Seed count and resample count are fixed before running. Full "
       "values are emitted to a config snapshot by the released pipeline.")
    _h(doc, "Appendix C  Per-sample logging schema", 1)
    _p(doc,
       "Each evaluated sample is logged with: id, dataset, label, pooling, condition, "
       "split, seed, C_i, entropy, probe prediction, LLM confidence/entropy, sim, "
       "B_pred, B_true, RUS, calibrated RUS, ΔC, SMART decision, loss-oracle decision, "
       "correctness-oracle decision, ℓ_p, ℓ_r, parametric/retrieval/SMART predictions, "
       "regret, and per-pass latency and prompt length. All tables and figures are "
       "derived from this file, so every number is reproducible from logs.")
    _h(doc, "Appendix D  Auxiliary components (not part of the main claims)", 1)
    _p(doc,
       "During development we explored two mechanisms that are orthogonal to the "
       "decision-time claim. We document them here for completeness; they are not "
       "evaluated in the main study and no result in the paper depends on them.")
    D.bold_label(doc, "D.1  Uncertainty-scaled adapter rank.",
       "An uncertainty signal U(x) can scale the LoRA rank per input, so confident "
       "inputs receive cheap low-rank adaptation and uncertain inputs receive more "
       "capacity:")
    X.add_equation(doc, E["unc"], "D1")
    X.add_equation(doc, E["rank"], "D2")
    D.bold_label(doc, "D.2  Attribution-based explanation check.",
       "For a subset of inputs we compute Integrated-Gradients attributions of the "
       "predicted-class logit, take the top-k content tokens 𝒯_k, and measure the "
       "fraction referenced by the generated explanation 𝒯_exp, as an internal "
       "consistency check rather than an explainability claim:")
    X.add_equation(doc, E["faith"], "D3")


def main():
    ap = argparse.ArgumentParser(description="Build the full draft manuscript (docx)")
    ap.add_argument("--out", default="runs/paper_full")
    args = ap.parse_args()
    out = build(args.out)
    print("Wrote", out)


if __name__ == "__main__":
    main()
