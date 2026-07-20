"""python-docx helpers for building the manuscript and supplementary documents."""
from __future__ import annotations

from pathlib import Path
from typing import Optional

import pandas as pd

TBD = "[[TBD-from-run]]"


def new_document():
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    return doc


def title(doc, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    return p


def centered(doc, text, italic=False, size=11):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    return p


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text):
    return doc.add_paragraph(text)


def bold_label(doc, label, text):
    p = doc.add_paragraph()
    p.add_run(label + " ").bold = True
    p.add_run(text)
    return p


def equation(doc, text, number: Optional[str] = None):
    """Render an equation as a centered monospaced line (draft typesetting)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10.5)
    if number:
        p.add_run("      (" + number + ")")
    return p


def _fmt_cell(v):
    if isinstance(v, float):
        if v != v:  # NaN == an undefined metric (e.g. precision with no positives)
            return "n/a"
        return f"{v:.3f}"
    return str(v)


def table_from_df(doc, df: pd.DataFrame, caption: str,
                  caption_above: bool = True):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if caption_above:
        cap = doc.add_paragraph()
        cap.add_run(caption).bold = True
    tbl = doc.add_table(rows=1, cols=len(df.columns))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].paragraphs[0].add_run(str(col)).bold = True
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for j, col in enumerate(df.columns):
            cells[j].text = _fmt_cell(row[col])
    if not caption_above:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(caption).italic = True
    doc.add_paragraph()
    return tbl


def figure(doc, image_path: str, caption: str, width_in: float = 5.5):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    if not Path(image_path).exists():
        p = doc.add_paragraph()
        p.add_run(f"[figure missing: {image_path} — run analysis on GPU box]").italic = True
        return p
    doc.add_picture(image_path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption).italic = True
    doc.add_paragraph()


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")
