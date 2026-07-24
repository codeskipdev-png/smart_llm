"""Generate the SMART-LLM PRE-EXPERIMENTAL PREDICTION manuscript (.docx).

This is a *scientific forecast* document, NOT an experimental paper. No
experiments have been conducted. Every quantitative value is an AI-predicted
experimental outcome (before execution), carried as a point estimate + range with
a per-table confidence level and reasoning. A red banner and per-table markers
make this explicit so the file can never be mistaken for measured results.

Real Word equations are rendered via OMML (editable in Word). Self-contained:
depends only on python-docx. Run:
    python -m smart_llm.paper.make_prediction_paper [--out DIR]
"""
from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd

from . import docx_utils as D
from . import omml as X

PRED = "AI-predicted experimental outcome (before execution)"


# ======================================================================== #
# Predicted tables (forecasts — NOT measurements). String cells preserve
# point/range formatting; every caption is marked as a prediction.
# ======================================================================== #
T1 = pd.DataFrame([
    ["No Retrieval", "0.60 / 0.57–0.63", "0.62", "0.58", "0.59", "0.00"],
    ["Always RAG", "0.72 / 0.69–0.75", "0.73", "0.71", "0.72", "1.00"],
    ["Confidence-based", "0.67 / 0.64–0.70", "0.69", "0.65", "0.67", "~0.55"],
    ["Entropy-based", "0.67 / 0.63–0.70", "0.68", "0.65", "0.66", "~0.58"],
    ["Adaptive-RAG-style", "0.68 / 0.65–0.71", "0.70", "0.66", "0.68", "~0.50"],
    ["SMART-LLM (ours)", "0.685 / 0.66–0.71", "0.71", "0.67", "0.68", "~0.33"],
], columns=["System", "Accuracy (pt / range)", "Macro-P", "Macro-R",
            "Macro-F1", "Retrieval rate"])

T2 = pd.DataFrame([
    ["SMART-LLM (last)", "0.68 / 0.64–0.72", "0.72", "0.62", "0.67", "0.15 / 0.10–0.22"],
    ["SMART-LLM (mean)", "0.64", "0.70", "0.57", "0.63", "0.18"],
    ["SMART-LLM (attention)", "0.66", "0.73", "0.58", "0.65", "0.16"],
    ["Oracle (upper bound)", "1.00", "1.00", "1.00", "1.00", "0.00"],
], columns=["Method (pooling)", "Agreement", "Precision", "Recall", "F1",
            "Mean regret (CE)"])

T3 = pd.DataFrame([
    ["B_true (bounded)", "0.30 / 0.20–0.40", "0.42 / 0.35–0.50", "0.55 / 0.45–0.65"],
], columns=["Target", "R^2", "MAE", "Pearson r"])

T4 = pd.DataFrame([
    ["SMART-LLM (full)", "0.68", "0.67", "0.15", "0.685", "0.33"],
    ["− RBE (similarity-only)", "0.66", "0.65", "0.17", "0.675", "0.34"],
    ["− Calibration (raw RUS)", "0.62", "0.63", "0.20", "0.660", "0.48"],
    ["Random (budget-matched)", "0.52", "0.36", "0.30", "0.640", "0.33"],
], columns=["Variant", "Agreement", "F1", "Mean regret (CE)", "Accuracy",
            "Retrieval rate"])

T5 = pd.DataFrame([
    ["Raw LLM verbalizer confidence", "0.30 / 0.22–0.38", "0.34 / 0.28–0.40"],
    ["Calibrated probe C_i", "0.06 / 0.04–0.09", "0.17 / 0.14–0.20"],
], columns=["Confidence signal", "ECE", "Brier"])

T6 = pd.DataFrame([
    ["No Retrieval", "0.00", "~350 / 300–420", "~70 / 55–95", "0.35"],
    ["Always RAG", "1.00", "~1400 / 1200–1600", "~180 / 150–210", "1.00"],
    ["SMART-LLM (ours)", "~0.33", "~700 / 600–820", "~120 / 100–140", "~0.65"],
], columns=["System", "Retrieval rate", "Avg prompt tokens",
            "Latency (ms/sample)", "Rel. compute"])

T7 = pd.DataFrame([
    ["20 Newsgroups (topic)", "0.30", "0.55", "0.68", "+0.02 (pred. n.s.)",
     "0.685", "0.72", "0.60"],
    ["Financial PhraseBank (sentiment)", "0.40", "0.62", "0.70", "+0.06 (pred. sig.)",
     "0.80", "0.81", "0.78"],
], columns=["Dataset", "RBE R^2", "Pearson r", "Router agreement",
            "Δ Agreement (full − sim-only)", "SMART acc", "Always-RAG acc",
            "No-retr acc"])


# ======================================================================== #
# Equation library (OMML)
# ======================================================================== #
def _eqs():
    L = X
    Ci = L.sub(L.var("C"), L.var("i"))
    hL = L.sub(L.var("h"), L.var("L"))
    muN = L.sub(L.var("μ"), L.var("𝒩"))
    lp = L.sub(L.var("ℓ"), L.var("p"))
    lr = L.sub(L.var("ℓ"), L.var("r"))
    Bpred = L.sub(L.var("B"), L.txt("pred"))
    Btrue = L.sub(L.var("B"), L.txt("true"))
    return {
        "benefit": L.concat(L.var("b"), L.delim(L.var("x")), L.op("="),
                            lp, L.op("−"), lr),
        "oracle": L.concat(
            L.sup(L.var("π"), L.op("⋆")), L.delim(L.var("x")), L.op("="),
            L.indicator(L.concat(lr, L.op("<"), lp)), L.op("="),
            L.indicator(L.concat(L.var("b"), L.delim(L.var("x")), L.op(">"), L.op("0")))),
        "regret": L.concat(
            L.var("R"), L.delim(L.var("π")), L.op("="),
            L.sub(L.op("𝔼"), L.var("x")), L.brack(L.concat(
                L.absv(L.concat(L.var("b"), L.delim(L.var("x")))), L.op("⋅"),
                L.indicator(L.concat(L.var("π"), L.delim(L.var("x")), L.op("≠"),
                            L.sup(L.var("π"), L.op("⋆")), L.delim(L.var("x"))))))),
        "conf": L.concat(
            Ci, L.op("="), L.under("max", L.var("j")),
            L.sub(L.func("softmax", L.concat(L.sub(L.var("W"), L.var("p")), hL)),
                  L.var("j"))),
        "rbe": L.concat(Bpred, L.op("="), L.func("RBE",
                        L.brack(L.concat(hL, L.op(" ; "), muN)))),
        "btrue": L.concat(
            Btrue, L.op("="), L.func("clip", L.concat(
                L.frac(L.concat(lp, L.op("−"), lr),
                       L.concat(L.absv(lp), L.op("+"), L.var("τ"))),
                L.op(", "), L.op("−"), L.var("c"), L.op(", "), L.var("c")))),
        "rus": L.concat(L.txt("RUS"), L.delim(L.concat(L.var("x"), L.op(", "),
                        L.var("𝒩"))), L.op("="), L.var("α"), L.op("⋅"),
                        L.func("sim", L.concat(L.var("x"), L.op(", "), L.var("𝒩"))),
                        L.op("+"), L.var("β"), L.op("⋅"), Bpred),
        "delta": L.concat(L.var("Δ"), L.var("C"), L.delim(L.var("x")), L.op("="),
                          L.func("cal", L.txt("RUS")), L.op("−"), Ci),
        "decide": L.concat(L.func("retrieve", L.var("x")), L.op("="),
                           L.indicator(L.concat(L.var("Δ"), L.var("C"),
                                       L.delim(L.var("x")), L.op(">"), L.op("0")))),
        "unc": L.concat(L.var("U"), L.delim(L.var("x")), L.op("="),
                        L.var("λ"), L.op("⋅"), L.sub(L.var("H"), L.txt("norm")),
                        L.op("+"), L.delim(L.concat(L.op("1"), L.op("−"), L.var("λ"))),
                        L.delim(L.concat(L.op("1"), L.op("−"), Ci))),
        "rank": L.concat(L.var("r"), L.delim(L.var("x")), L.op("="),
                         L.sub(L.var("r"), L.txt("min")), L.op("+"),
                         L.delim(L.concat(L.sub(L.var("r"), L.txt("max")), L.op("−"),
                                 L.sub(L.var("r"), L.txt("min")))), L.op("⋅"),
                         L.var("U"), L.delim(L.var("x"))),
    }


def _p(doc, t): return D.para(doc, t)
def _h(doc, t, l=1): return D.heading(doc, t, l)


def _pred_table(doc, df, cap, confidence, reasoning):
    """A predicted-results table with its mandatory confidence + reasoning."""
    D.table_from_df(doc, df, cap + f"  [{PRED}]")
    D.bold_label(doc, "Prediction confidence:", confidence)
    D.bold_label(doc, "Reasoning:", reasoning)
    doc.add_paragraph()


# ======================================================================== #
# Manuscript
# ======================================================================== #
def build(out_dir: str = "runs/paper_full") -> str:
    from docx.shared import RGBColor
    E = _eqs()
    doc = D.new_document()

    D.title(doc, "Confidence-Aware Retrieval and Mixture-of-LoRA Adaptation for "
                 "Explainable Few-Shot Large Language Model Text Classification "
                 "(SMART-LLM)")
    D.centered(doc, "A Pre-Experimental Prediction Manuscript (Scientific Forecast)",
               italic=True, size=12)
    D.centered(doc, "Anonymous Author(s)")
    D.centered(doc, "Prepared for Information Sciences / Knowledge-Based Systems / "
                    "IEEE T-AI", italic=True, size=10)

    # prediction banner (red)
    b = doc.add_paragraph()
    r = b.add_run(
        "STATUS — READ FIRST. This is a PRE-EXPERIMENTAL PREDICTION manuscript, not "
        "an experimental paper. NO EXPERIMENTS HAVE BEEN CONDUCTED. Every quantitative "
        "value in Sections 5 and 7 is an AI-PREDICTED EXPERIMENTAL OUTCOME (BEFORE "
        "EXECUTION): a forecast with an uncertainty range and a per-table confidence "
        "level, to be compared against real RTX 4090 results later. These numbers must "
        "not be cited as measurements.")
    r.italic = True; r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    doc.add_paragraph()

    _abstract(doc)
    _introduction(doc)
    _related(doc)
    _methodology(doc, E)
    _setup(doc)
    _results(doc)
    _analysis(doc)
    _cases(doc)
    _discussion(doc)
    _conclusion(doc)
    _references(doc)

    Path(out_dir).mkdir(parents=True, exist_ok=True)
    out = str(Path(out_dir) / "SMART_LLM_prediction.docx")
    doc.save(out)
    return out


def _abstract(doc):
    _h(doc, "Abstract", 1)
    _p(doc,
       "Retrieval augmentation improves few-shot large language model (LLM) text "
       "classification on many inputs but not all: when the parametric model is already "
       "correct, retrieval adds cost, and when retrieved context is off-topic or "
       "misleading, it can displace a correct prediction. We study whether an LLM-based "
       "system can decide, before retrieving, whether retrieval will help a given "
       "input. SMART-LLM combines a calibrated internal-confidence probe over a frozen "
       "LLM, a Retrieval Benefit Estimator (RBE) that predicts the expected reduction "
       "in classification loss from retrieval using only pre-retrieval features, a "
       "Retrieval Utility Score (RUS) fusing similarity and predicted benefit, "
       "confidence calibration that places utility and confidence on a common scale, "
       "and an auxiliary Mixture-of-LoRA adaptation that scales adapter capacity with "
       "input uncertainty. Because the decision uses only cached pre-retrieval "
       "features, no retrieval-augmented forward pass is spent to decide. This "
       "manuscript predicts, prior to execution, the likely outcomes of a planned "
       "evaluation on 20 Newsgroups (primary) and Financial PhraseBank (secondary) with "
       "a frozen Qwen2.5-7B-Instruct backbone and bge-large embeddings, against six "
       "baselines. We forecast that selective arbitration recovers most of the accuracy "
       "benefit of always-on retrieval at roughly one-third of the retrieval rate, that "
       "calibration is the single most important component, that the learned benefit "
       "term adds measurable value chiefly where semantic similarity is a weak cue "
       "(sentiment more than topic), and that the clearest advantage over always-on "
       "retrieval appears under corrupted retrieval. We state confidence levels and "
       "falsification conditions for each prediction. All results herein are "
       "AI-predicted experimental outcomes (before execution).")
    D.bold_label(doc, "Keywords:",
       "decision-time retrieval; selective prediction; adaptive computation; retrieval "
       "benefit estimation; confidence calibration; regret analysis; parameter-"
       "efficient fine-tuning; few-shot text classification; large language models; "
       "pre-registered prediction.")


def _introduction(doc):
    _h(doc, "1  Introduction", 1)
    _p(doc,
       "Augmenting a frozen LLM with retrieved in-context examples reliably improves "
       "many, but not all, few-shot text-classification inputs. In its default form the "
       "augmentation is unconditional: every input triggers a retrieval and a longer, "
       "more expensive forward pass. This is wasteful when the parametric model already "
       "answers correctly and harmful when the retrieved context is off-topic, "
       "redundant, or adversarial and displaces a correct parametric prediction. The "
       "quantity that matters — whether retrieval will help this input — is normally "
       "observed only after retrieving and running the augmented model, which is "
       "precisely the cost one would like to avoid.")
    _p(doc, "We treat retrieval as a per-input decision under uncertainty and ask a "
            "single question:")
    D.bold_label(doc, "Research question.",
       "Can an LLM-based system estimate, before performing retrieval, whether "
       "retrieval will improve its prediction for a given input, and route on that "
       "estimate to preserve accuracy while avoiding unnecessary and harmful retrieval?")
    _p(doc,
       "This reframing places the problem alongside selective prediction and adaptive "
       "computation, rather than alongside methods that improve retrieval quality "
       "itself. SMART-LLM instantiates it with two pre-retrieval signals — a calibrated "
       "internal confidence and a predicted retrieval benefit — combined by a "
       "calibrated comparison. Two auxiliary mechanisms (Mixture-of-LoRA adaptation and "
       "attribution-based explanation verification) are included for completeness and "
       "are not central to the decision-time claim.")
    _p(doc, "We frame the contributions as claims to be tested, not yet verified:")
    D.numbered(doc, "A formulation of the retrieval decision as pre-retrieval benefit "
                    "prediction, learned from an offline ground-truth benefit signal.")
    D.numbered(doc, "A calibrated arbitration rule fusing confidence and utility, "
                    "analysed by agreement with an oracle policy and by regret.")
    D.numbered(doc, "A pre-experimental forecast, with explicit confidence and "
                    "falsification criteria, of how the method and six baselines will "
                    "behave — the object of this manuscript.")
    _p(doc,
       "Because this is a forecast, the Introduction makes no empirical claim; Section 5 "
       "states predictions and Section 9 states what would falsify them.")


def _related(doc):
    _h(doc, "2  Related Work", 1)
    D.bold_label(doc, "Selective prediction and decision-theoretic inference.",
       "Selective classifiers abstain when confidence is low, trading coverage for "
       "reliability (Geifman & El-Yaniv, 2017; El-Yaniv & Wiener, 2010). We adopt the "
       "same per-input framing — 'retrieve vs. trust the parametric model' in place of "
       "'predict vs. abstain' — and inherit its evaluation discipline (oracle "
       "comparison, regret).")
    D.bold_label(doc, "Adaptive computation.",
       "Early-exit networks (Xin et al., 2020), conditional computation and mixtures of "
       "experts (Fedus et al., 2022), and confident adaptive language modelling "
       "(Schuster et al., 2022) allocate internal compute per input. SMART-LLM shares "
       "the objective but targets an external resource — whether to spend a "
       "retrieval-augmented pass — and decides before the expensive pass from cached "
       "features.")
    D.bold_label(doc, "Adaptive and selective retrieval (closest work).",
       "RAG conditions predictions on retrieved evidence (Lewis et al., 2020). One "
       "family decides during decoding: Self-RAG emits retrieval-control tokens (Asai "
       "et al., 2024); FLARE triggers on next-token uncertainty (Jiang et al., 2023). A "
       "second family decides before generation but from coarse signals: popularity "
       "gating (Mallen et al., 2023), learned query-complexity routing (Adaptive-RAG; "
       "Jeong et al., 2024), and self-knowledge elicitation (SKR; Wang et al., 2023), "
       "building on the finding that LLMs are partly aware of their own competence "
       "(Kadavath et al., 2022). SMART-LLM belongs to the decide-before-retrieving "
       "family but regresses a continuous, loss-calibrated benefit rather than "
       "classifying complexity/popularity, reads the hidden state and neighbour "
       "centroid cached before any augmented pass, and scores against an oracle by "
       "agreement and regret. We do not claim to beat these methods on their native QA "
       "benchmarks; we isolate the pre-retrieval benefit-estimation question and "
       "include an entropy-gated (FLARE/Adaptive-RAG-style) baseline.")
    D.bold_label(doc, "Calibration and PEFT.",
       "Temperature/Platt/isotonic calibration (Guo et al., 2017; Platt, 1999; "
       "Zadrozny & Elkan, 2002) align probabilities with accuracy; we use calibration "
       "as a control signal so utility and confidence are comparable. LoRA (Hu et al., "
       "2022) and Integrated Gradients (Sundararajan et al., 2017) support the "
       "auxiliary components only.")


def _methodology(doc, E):
    _h(doc, "3  Methodology", 1)
    _h(doc, "3.1  Problem Definition", 2)
    _p(doc,
       "Let x be an input with gold label y over K classes. A frozen classifier incurs "
       "cross-entropy loss ℓ_p(x) without retrieval and ℓ_r(x) with retrieval of a "
       "neighbour set 𝒩 (|𝒩| = k; K denotes the number of classes and 𝒩 the retrieved "
       "set). The per-input retrieval benefit and the oracle policy are")
    X.add_equation(doc, E["benefit"], "1")
    X.add_equation(doc, E["oracle"], "2")
    _p(doc, "and any policy is measured by regret, the excess loss over the "
            "pointwise-optimal choice:")
    X.add_equation(doc, E["regret"], "3")
    D.bold_label(doc, "Loss vs. correctness.",
       "b(x) is defined on cross-entropy, so the oracle above is a loss oracle; a "
       "retrieval that lowers loss need not flip the arg-max label. We therefore also "
       "track a correctness oracle (retrieval changes an incorrect parametric label to "
       "correct) and report which quantity each table measures.")
    _h(doc, "3.2  SMART-LLM Architecture", 2)
    _p(doc,
       "A frozen instruction-tuned LLM classifies via a letter-verbalizer (each class "
       "maps to an option token), so one forward pass yields a class distribution, a "
       "cross-entropy loss, and pooled final-layer hidden states h_L. The arbiter "
       "consumes h_L, a calibrated confidence C_i, the retrieved-neighbour centroid "
       "μ_𝒩, the query–neighbour similarity sim, and a predicted benefit B_pred. All "
       "are available before the augmented pass, which runs only for selected inputs.")
    _h(doc, "3.3  Retrieval Benefit Estimator", 2)
    _p(doc, "The RBE is a small MLP predicting the loss reduction from retrieval using "
            "pre-retrieval features only:")
    X.add_equation(doc, E["rbe"], "4")
    _p(doc, "It is trained against a numerically stable ground-truth benefit obtained "
            "offline by running the frozen LLM with and without retrieval (used only "
            "to build supervision, never to decide):")
    X.add_equation(doc, E["btrue"], "5")
    _p(doc, "with a denominator floor τ (avoiding blow-up as ℓ_p → 0) and clip c, "
            "trained with a robust Huber objective. The floor/clip change magnitudes "
            "but preserve the sign, so the oracle decision is unaffected.")
    _h(doc, "3.4  Retrieval Utility Score", 2)
    _p(doc, "RUS fuses semantic similarity and predicted benefit, with α, β tuned on "
            "validation by oracle agreement:")
    X.add_equation(doc, E["rus"], "6")
    _h(doc, "3.5  Confidence Calibration", 2)
    _p(doc,
       "The raw verbalizer confidence of instruction-tuned LLMs is typically "
       "over-confident. A lightweight probe with temperature-scaled logits gives")
    X.add_equation(doc, E["conf"], "7")
    _p(doc, "RUS is mapped to a probability scale by a calibrator fitted on a held-out "
            "split, and the arbiter compares the two calibrated quantities:")
    X.add_equation(doc, E["delta"], "8")
    X.add_equation(doc, E["decide"], "9")
    _p(doc,
       "Interpreting cal(RUS) as an estimate of P(b(x) > 0 | x) and C_i as an estimate "
       "of the probability the parametric prediction is correct, the rule retrieves "
       "when the estimated probability that retrieval helps exceeds the estimated "
       "probability that the internal answer is already correct.")
    _h(doc, "3.6  Mixture-of-LoRA Adaptation (auxiliary)", 2)
    _p(doc, "As a complementary efficiency mechanism, an uncertainty signal selects "
            "among a small bank of LoRA experts / scales adapter rank per input:")
    X.add_equation(doc, E["unc"], "10")
    X.add_equation(doc, E["rank"], "11")
    _p(doc, "so confident inputs receive cheap low-rank adaptation and uncertain inputs "
            "receive more capacity. This is not central to the decision-time claim and "
            "is evaluated separately.")
    _h(doc, "3.7  Inference Algorithm", 2)
    for line in [
        "1. Parametric pass (frozen LLM): obtain class distribution, ℓ_p, h_L, "
        "calibrated C_i.",
        "2. Retrieve k neighbours 𝒩; compute sim and centroid μ_𝒩 (vectors only, no "
        "LLM pass).",
        "3. B_pred = RBE([h_L; μ_𝒩]); RUS = α·sim + β·B_pred; compute cal(RUS).",
        "4. ΔC = cal(RUS) − C_i.",
        "5. If ΔC > 0: run the retrieval-augmented pass (only now is it spent) → "
        "prediction, ℓ_r; else return the parametric prediction.",
        "6. (auxiliary) select LoRA capacity from U(x); (auxiliary) attribution-based "
        "explanation check.",
    ]:
        D.bullet(doc, line)


def _setup(doc):
    _h(doc, "4  Experimental Setup", 1)
    _h(doc, "4.1  Datasets", 2)
    _p(doc,
       "Primary — 20 Newsgroups: 20 topical classes, several fine-grained and "
       "overlapping (e.g., comp.sys.ibm.pc.hardware vs. comp.sys.mac.hardware; the "
       "politics/religion clusters). Headers/footers/quotes removed; stratified "
       "evaluation subset. Secondary — Financial PhraseBank: 3-class sentence-level "
       "sentiment, used to test generalization to a regime where semantic similarity is "
       "a weaker cue for retrieval usefulness.")
    _h(doc, "4.2  Implementation Details", 2)
    _p(doc,
       "Frozen Qwen2.5-7B-Instruct (bf16) on a single RTX 4090 (24 GB); embeddings "
       "bge-large with a FAISS inner-product index over the training pool; k = 8 "
       "neighbours. Only the confidence probe, the RBE, the calibration map, and "
       "(auxiliary) LoRA parameters are trained; the LLM is frozen and its forward "
       "passes are cached so it runs once. Last-token pooling selected by validation "
       "routing agreement. Planned statistical protocol: 5 seeds; 95% bootstrap "
       "confidence intervals (10,000 resamples); paired McNemar (accuracy) and paired "
       "bootstrap (regret/agreement); differences significant only at p < 0.05.")
    _h(doc, "4.3  Baselines", 2)
    _p(doc,
       "(1) No Retrieval; (2) Always Retrieval-Augmented Generation; (3) "
       "Confidence-based retrieval (retrieve iff C_i < τ); (4) Entropy-based retrieval "
       "(retrieve iff predictive entropy > τ); (5) Adaptive-RAG-style routing "
       "(threshold/learned complexity gate); (6) SMART-LLM. A budget-matched Random "
       "policy is included in the ablation to separate 'decide well' from 'retrieve "
       "less'.")
    _h(doc, "4.4  Evaluation Metrics", 2)
    _p(doc,
       "Accuracy, macro precision/recall/F1; retrieval rate; latency and token cost; "
       "calibration (ECE, Brier); router quality vs. oracle (agreement, "
       "precision/recall/F1, regret); RBE quality (R², MAE, Pearson r).")


def _results(doc):
    _h(doc, "5  Predicted Experimental Results", 1)
    _p(doc,
       "All values below are AI-predicted experimental outcomes (before execution). "
       "Each is a point forecast with a plausible range; each table carries a "
       "confidence level (High / Medium / Low) and reasoning. Ordering/direction "
       "predictions are generally more reliable than absolute magnitudes (Section 9). "
       "Regret is reported in cross-entropy-loss units (excess CE over the oracle "
       "choice).")

    _h(doc, "5.1  Table 1 — Main performance (20 Newsgroups, clean retrieval)", 2)
    _pred_table(doc, T1,
       "Table 1. Predicted main performance (20 Newsgroups, clean retrieval, test "
       "split).",
       "Medium (relative ordering: High; absolute accuracy magnitudes: Medium).",
       "20 Newsgroups is a topical task where retrieved same-topic exemplars usually "
       "help, so Always-RAG is predicted to be the clean-accuracy ceiling. SMART-LLM is "
       "predicted to sit just below it while retrieving for ~1/3 of inputs, and to "
       "match or slightly exceed the other selective baselines at a lower retrieval "
       "rate — a better accuracy/retrieval trade-off rather than a higher accuracy "
       "peak. We deliberately do not predict SMART-LLM > Always-RAG on clean accuracy; "
       "claiming so would be inconsistent with the mechanism.")

    _h(doc, "5.2  Table 2 — Router vs. oracle (20 Newsgroups)", 2)
    _pred_table(doc, T2,
       "Table 2. Predicted router vs. oracle by pooling.",
       "Medium.",
       "Agreement is predicted well above the 0.5 chance level but clearly below 1.0, "
       "because benefit sign is only partly recoverable from pre-retrieval features. "
       "Last-token pooling is predicted strongest. Precision > recall is predicted: the "
       "calibrated rule is conservative, retrieving mainly where utility clearly "
       "exceeds confidence, so it misses some beneficial retrievals (lower recall) "
       "while being comparatively accurate when it does retrieve.")

    _h(doc, "5.3  Table 3 — RBE prediction quality (20 Newsgroups)", 2)
    _pred_table(doc, T3,
       "Table 3. Predicted RBE prediction quality on the bounded benefit target.",
       "Low–Medium (among the least certain predictions).",
       "h_L encodes parametric difficulty and μ_𝒩 encodes retrieval content, but their "
       "interaction — whether this context helps this input — is only partly linearly "
       "recoverable. We predict a positive but moderate R² and correlation, and "
       "emphasise that the RBE is expected to be most reliable at the sign level (which "
       "is what routing needs), not as a precise magnitude regressor. R² is sensitive "
       "to the benefit distribution and the floor/clip, so its realized value could "
       "deviate substantially.")

    _h(doc, "5.4  Table 4 — Ablation (20 Newsgroups)", 2)
    _pred_table(doc, T4,
       "Table 4. Predicted ablation: router quality and accuracy.",
       "Medium (direction of each effect: High; magnitudes: Medium).",
       "Three directional predictions are made with high confidence. Removing "
       "calibration is predicted the most damaging change — agreement falls, regret "
       "rises, retrieval rate inflates (an uncalibrated utility over-fires). Removing "
       "the RBE on a topical dataset is predicted to cost little (similarity already "
       "proxies usefulness for topic); the Δ over similarity-only routing is predicted "
       "small and possibly not significant on 20NG. The budget-matched Random policy is "
       "predicted well below all learned routers, isolating that the value is in "
       "deciding well, not merely in retrieving less.")

    _h(doc, "5.5  Table 5 — Calibration (20 Newsgroups)", 2)
    _pred_table(doc, T5,
       "Table 5. Predicted calibration of confidence signals.",
       "High.",
       "Temperature/Platt scaling reliably reduces the systematic over-confidence of "
       "instruction-tuned LLM verbalizer probabilities; a large ECE/Brier improvement "
       "is one of the most robust findings in the calibration literature. This is our "
       "highest-confidence prediction and underwrites the calibration ablation in "
       "Table 4.")

    _h(doc, "5.6  Table 6 — Computational efficiency (20 Newsgroups)", 2)
    _pred_table(doc, T6,
       "Table 6. Predicted computation: retrieval rate, prompt length, latency, "
       "relative compute.",
       "High for ordering and token cost; Medium for absolute latency.",
       "Prompt length is dominated by retrieved demonstrations, so token cost tracks "
       "retrieval rate almost mechanically; the ordering No-Retrieval < SMART-LLM < "
       "Always-RAG is near-certain. Absolute latency depends on batching, KV-cache "
       "reuse, and prefill efficiency on the RTX 4090, so the millisecond values carry "
       "more uncertainty. We explicitly predict SMART-LLM is not a uniform latency win "
       "over No-Retrieval (it always pays one parametric pass); its advantage is fewer "
       "augmented passes and shorter prompts than Always-RAG, plus robustness.")

    _h(doc, "5.7  Table 7 — Cross-dataset generalization", 2)
    _pred_table(doc, T7,
       "Table 7. Predicted cross-dataset generalization (20 Newsgroups vs. Financial "
       "PhraseBank).",
       "Medium (topic row); Low–Medium (the cross-dataset Δ Agreement claim is the "
       "single riskiest prediction).",
       "The central mechanistic prediction is that the learned benefit term earns its "
       "place where similarity is a weak cue. Sentiment polarity is less aligned with "
       "embedding similarity than topic, so we predict the RBE's marginal gain over "
       "similarity-only routing is larger and (unlike 20NG) statistically significant "
       "on Financial PhraseBank, with a higher RBE R². We also predict retrieval helps "
       "less overall on sentiment (No-Retrieval already high), so all three systems "
       "cluster. This row is where real results are most likely to diverge.")


def _analysis(doc):
    _h(doc, "6  Analysis", 1)
    D.bold_label(doc, "Why retrieval sometimes hurts.",
       "Retrieved demonstrations enter the prompt as evidence the model attends to. "
       "When neighbours share the true label they reinforce the correct class; when "
       "they are topically similar but wrong-labelled, redundant, or adversarial, they "
       "can shift probability mass toward an incorrect class and override a correct "
       "parametric prediction. The expected harm grows as retrieval quality degrades "
       "and as the model's own competence on the input rises. We therefore predict "
       "always-on retrieval falls below the no-retrieval baseline under a constructed "
       "adversarial condition, while a selective policy that declines low-utility "
       "retrievals will not.")
    D.bold_label(doc, "Why learned arbitration should outperform similarity-only "
                      "routing — and when it should not.",
       "Similarity measures whether neighbours are close, not whether they are useful "
       "for this decision. On topic classification the two coincide strongly, so we "
       "predict the learned RBE adds little over similarity-only routing on 20 "
       "Newsgroups. On sentiment, closeness is a weaker proxy for label-usefulness, so "
       "we predict the RBE's model-internal signal (via h_L) contributes information "
       "similarity alone lacks. This is a conditional, testable prediction, not a "
       "universal claim.")
    D.bold_label(doc, "Failure cases (anticipated).",
       "(i) Confidently wrong parametric answers — high C_i but incorrect — suppress "
       "beneficial retrieval, bounded by calibration quality but not eliminated. (ii) "
       "Misleading similar neighbours inflate RUS and trigger harmful retrieval. (iii) "
       "RBE mis-estimation on out-of-distribution inputs mis-routes. (iv) Calibration "
       "drift under distribution shift decouples C_i from true correctness. These are "
       "the structural failure modes of any confidence-plus-utility rule.")
    D.bold_label(doc, "Expected limitations.",
       "Two datasets and two task types do not establish general regimes; the method "
       "always pays one parametric pass; benefit magnitude is predicted only partly "
       "learnable; and the auxiliary Mixture-of-LoRA and explanation-verification "
       "components are not expected to carry the core result.")


def _cases(doc):
    _h(doc, "7  Case Studies", 1)
    _p(doc,
       "The following ten cases describe predicted qualitative behavior (before "
       "execution); they illustrate the mechanism and are not transcripts of real runs.")
    _h(doc, "7.1  Predicted successful cases", 2)
    for t in [
        "Ambiguous cross-posted 20NG document (between comp.graphics and "
        "comp.os.ms-windows.misc): parametric confidence low, neighbour similarity high "
        "and same-label → cal(RUS) > C_i → SMART retrieves → predicted correct. "
        "Uncertainty opens the gate; useful neighbours cross it.",
        "Clear-cut rec.sport.baseball post: high C_i, low predicted benefit → ΔC < 0 → "
        "SMART skips retrieval → predicted correct at reduced compute. Confidence vetoes "
        "unnecessary retrieval.",
        "Negated financial sentence ('results were not as weak as feared'): parametric "
        "model wavers between neutral/positive; similar labelled exemplars disambiguate "
        "→ SMART retrieves → predicted correct. Benefit predicted where polarity cues "
        "are subtle.",
        "Adversarially corrupted retrieval: injected hard negatives lower sim and "
        "predicted utility → RUS low → SMART declines retrieval → predicted correct "
        "where Always-RAG is misled. The safeguard always-on policies lack.",
        "Rare-class financial input with weak neighbour similarity but informative "
        "hidden state: sim low yet RBE predicts benefit from h_L → SMART retrieves → "
        "predicted correct. The learned term acting beyond similarity.",
    ]:
        D.numbered(doc, t)
    _h(doc, "7.2  Predicted failure cases", 2)
    for t in [
        "Confidently wrong parametric answer (high C_i, incorrect): ΔC < 0 suppresses "
        "beneficial retrieval → predicted wrong. Bounded by calibration but not removed.",
        "Topically similar but misleading neighbour: high sim inflates RUS → SMART "
        "retrieves → context corrupts an otherwise-correct parametric prediction → "
        "predicted wrong.",
        "Out-of-distribution input for the RBE: benefit mis-estimated → mis-routing in "
        "either direction → predicted wrong.",
        "Genuinely ambiguous multi-topic document: neither parametric nor retrieval "
        "path is correct → predicted wrong regardless of routing (a data-limited, not "
        "method-limited, failure).",
        "Calibration drift under domain shift (probe fit on one corpus, applied to "
        "another): C_i decouples from true correctness → systematic routing errors → "
        "predicted wrong on a subset.",
    ]:
        D.numbered(doc, t)


def _discussion(doc):
    _h(doc, "8  Discussion", 1)
    D.bold_label(doc, "Scientific contribution (as proposed).",
       "The framing of the retrieval decision as pre-retrieval benefit prediction "
       "scored against an oracle by agreement and regret, and the explicit dependence "
       "of routing quality on calibration, are the intended contributions. If the "
       "predictions hold, the practical value is a favourable accuracy/compute "
       "trade-off and a robustness safeguard, not a new accuracy ceiling.")
    D.bold_label(doc, "Limitations.",
       "Narrow task coverage (topic + sentiment); dependence on retriever and embedding "
       "quality; the always-paid parametric pass; partial predictability of benefit "
       "magnitude; auxiliary components not central.")
    D.bold_label(doc, "Expected reviewer criticism (anticipated, with planned "
                      "responses).",
       "'The RBE adds little' — predicted true on topic; positioned as conditional and "
       "tested on sentiment (Table 7); if it fails there too we report the RBE as a "
       "negative result. 'The adversarial condition is a strawman' — agreed; presented "
       "as a constructed worst case, with the budget-matched Random comparison as the "
       "honest like-for-like evidence. 'Only two datasets; single backbone' — "
       "acknowledged as the primary threat to generality. 'The theory is idealised' — "
       "presented as design rationale, not a guarantee, with residuals measured "
       "empirically. 'Mixture-of-LoRA and explanation verification are underdeveloped' "
       "— scoped as auxiliary, with no core claim depending on them.")
    D.bold_label(doc, "How experiments may falsify the hypothesis.",
       "The central hypothesis is falsified if, on real data, a learned selective "
       "policy cannot beat a budget-matched random policy (agreement ≈ chance; "
       "accuracy-at-budget ≤ Random), if calibration does not improve routing, or if a "
       "single confidence threshold dominates the full method on the accuracy/compute "
       "frontier. These are explicit, pre-registered failure criteria.")


def _conclusion(doc):
    _h(doc, "9  Conclusion and Prediction Reliability Assessment", 1)
    _p(doc,
       "We have described SMART-LLM and, without conducting experiments, predicted the "
       "likely outcomes of a planned RTX 4090 evaluation, with per-table confidence and "
       "reasoning. The forecast's core is: selective arbitration recovers most of "
       "always-on retrieval's accuracy benefit at ~1/3 the retrieval rate; calibration "
       "is the most important component; the learned benefit term matters mainly where "
       "similarity is weak; and the clearest advantage over always-on retrieval appears "
       "under corrupted retrieval.")
    D.bold_label(doc, "1. Highly reliable predictions (High confidence).",
       "Large calibration improvement from temperature/Platt scaling (Table 5); compute/"
       "token ordering No-Retrieval < SMART-LLM < Always-RAG (Table 6); removing "
       "calibration degrades routing and inflates retrieval rate (Table 4); always-on "
       "retrieval degrades (and can fall below no-retrieval) under adversarial "
       "retrieval; learned routers beat a budget-matched Random policy (Table 4).")
    D.bold_label(doc, "2. Uncertain predictions (Medium–Low confidence).",
       "Absolute accuracy magnitudes for all systems (±5–8 absolute points plausible); "
       "RBE R² and Pearson r (Table 3); exact retrieval rates and the exact "
       "SMART-vs-Always-RAG accuracy gap; the cross-dataset Δ Agreement significance "
       "claim (Table 7); millisecond-level latency.")
    D.bold_label(doc, "3. Outcomes that would invalidate SMART-LLM.",
       "Router oracle agreement ≈ 0.5, or SMART accuracy-at-budget ≤ budget-matched "
       "Random; calibration failing to improve routing; RBE R² ≤ 0 and similarity-only "
       "≥ full on every dataset; a single confidence/entropy threshold Pareto-"
       "dominating the full method; always-on retrieval never being harmful in any "
       "realistic condition.")
    D.bold_label(doc, "4. Expected differences between prediction and real RTX 4090 "
                      "results.",
       "Directions and orderings are expected to hold more often than magnitudes. "
       "Realistic expectation: most High-confidence directional predictions hold; "
       "absolute accuracies within ~±5–8 points of the point forecasts; retrieval rates "
       "within ~±10–15 points; RBE R² potentially off by ±0.15; latency the most likely "
       "to differ. The riskiest items are Table 3 (RBE magnitude) and the Table 7 "
       "sentiment-regime claim.")
    D.bold_label(doc, "A necessary caution on expecting an exact match.",
       "A pre-experimental forecast that reproduced real results identically would be a "
       "statistical coincidence, not evidence of good reasoning — genuine measurements "
       "carry seed variance, implementation choices, and data idiosyncrasies that no "
       "prior forecast can hit to the decimal. The scientific success criterion is "
       "calibration, not identity: that the real numbers fall within the stated ranges, "
       "that the High-confidence directional claims hold, and that any pre-registered "
       "falsification condition either fails to trigger or, if it does, is reported "
       "honestly. Treating a non-identical outcome as a failure of the prediction would "
       "itself be a methodological error.")
    _p(doc,
       "Reminder: this is a scientific forecast manuscript. No experiments have been "
       "conducted. Every quantitative value is an AI-predicted experimental outcome "
       "(before execution), to be compared against real results and revised "
       "accordingly.")


def _references(doc):
    _h(doc, "References", 1)
    for r in [
        "Lewis, P., et al. (2020). Retrieval-Augmented Generation for Knowledge-"
        "Intensive NLP Tasks. NeurIPS.",
        "Asai, A., et al. (2024). Self-RAG: Learning to Retrieve, Generate, and "
        "Critique through Self-Reflection. ICLR.",
        "Jiang, Z., et al. (2023). Active Retrieval Augmented Generation (FLARE). EMNLP.",
        "Mallen, A., et al. (2023). When Not to Trust Language Models. ACL.",
        "Jeong, S., et al. (2024). Adaptive-RAG: Learning to Adapt Retrieval-Augmented "
        "LLMs through Question Complexity. NAACL.",
        "Wang, Y., et al. (2023). Self-Knowledge Guided Retrieval Augmentation (SKR). "
        "Findings of EMNLP.",
        "Kadavath, S., et al. (2022). Language Models (Mostly) Know What They Know. "
        "arXiv:2207.05221.",
        "Geifman, Y., & El-Yaniv, R. (2017). Selective Classification for Deep Neural "
        "Networks. NeurIPS.",
        "El-Yaniv, R., & Wiener, Y. (2010). On the Foundations of Noise-free Selective "
        "Classification. JMLR.",
        "Schuster, T., et al. (2022). Confident Adaptive Language Modeling (CALM). "
        "NeurIPS.",
        "Xin, J., et al. (2020). DeeBERT: Dynamic Early Exiting for Accelerating BERT "
        "Inference. ACL.",
        "Fedus, W., Zoph, B., & Shazeer, N. (2022). Switch Transformers. JMLR.",
        "Hu, E. J., et al. (2022). LoRA: Low-Rank Adaptation of Large Language Models. "
        "ICLR.",
        "Guo, C., et al. (2017). On Calibration of Modern Neural Networks. ICML.",
        "Platt, J. (1999). Probabilistic Outputs for Support Vector Machines.",
        "Zadrozny, B., & Elkan, C. (2002). Transforming Classifier Scores into Accurate "
        "Multiclass Probability Estimates. KDD.",
        "Sundararajan, M., Taly, A., & Yan, Q. (2017). Axiomatic Attribution for Deep "
        "Networks. ICML.",
        "Qwen Team (2024). Qwen2.5 Technical Report.",
        "Xiao, S., et al. (2023). C-Pack: Packed Resources for General Chinese and "
        "English Embeddings (BGE).",
        "Johnson, J., Douze, M., & Jégou, H. (2019). Billion-Scale Similarity Search "
        "with GPUs (FAISS). IEEE Transactions on Big Data.",
    ]:
        _p(doc, r)


def main():
    ap = argparse.ArgumentParser(description="Build the SMART-LLM prediction manuscript")
    ap.add_argument("--out", default="runs/paper_full")
    args = ap.parse_args()
    out = build(args.out)
    print("Wrote", out)


if __name__ == "__main__":
    main()
