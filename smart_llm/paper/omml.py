"""Office Math (OMML) builder for python-docx.

Produces *real* Word equations (editable in Word's equation editor), not plain
text. Equations are composed as OMML XML strings and inserted with parse_xml;
the whole equation is one string with a single ``xmlns:m`` declaration on the
root ``m:oMath`` so every child inherits the namespace.

Convention: variables render italic (math default); operators, function names,
and digits are marked upright via ``m:nor``.
"""
from __future__ import annotations

M = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _esc(t: str) -> str:
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ---- leaf runs ------------------------------------------------------------ #
def mr(text: str, upright: bool = False) -> str:
    rpr = "<m:rPr><m:nor/></m:rPr>" if upright else ""
    return f'<m:r>{rpr}<m:t xml:space="preserve">{_esc(text)}</m:t></m:r>'


def var(text: str) -> str:          # italic variable
    return mr(text, upright=False)


def op(text: str) -> str:           # upright operator / punctuation / digits
    return mr(text, upright=True)


def txt(text: str) -> str:          # upright words (function names, "if")
    return mr(text, upright=True)


# ---- structures ----------------------------------------------------------- #
def frac(num: str, den: str) -> str:
    return f"<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>"


def sub(base: str, s: str) -> str:
    return f"<m:sSub><m:e>{base}</m:e><m:sub>{s}</m:sub></m:sSub>"


def sup(base: str, s: str) -> str:
    return f"<m:sSup><m:e>{base}</m:e><m:sup>{s}</m:sup></m:sSup>"


def subsup(base: str, sb: str, sp: str) -> str:
    return (f"<m:sSubSup><m:e>{base}</m:e><m:sub>{sb}</m:sub>"
            f"<m:sup>{sp}</m:sup></m:sSubSup>")


def delim(inner: str, left: str = "(", right: str = ")") -> str:
    return (f'<m:d><m:dPr><m:begChr m:val="{left}"/><m:endChr m:val="{right}"/>'
            f"</m:dPr><m:e>{inner}</m:e></m:d>")


def absv(inner: str) -> str:
    return delim(inner, "|", "|")


def brack(inner: str) -> str:
    return delim(inner, "[", "]")


def sqrt(inner: str) -> str:
    return (f"<m:rad><m:radPr><m:degHide m:val=\"1\"/></m:radPr><m:deg/>"
            f"<m:e>{inner}</m:e></m:rad>")


def nary(chr_: str, lo: str, hi: str, body: str, limloc: str = "subSup") -> str:
    naryPr = (f'<m:naryPr><m:chr m:val="{chr_}"/><m:limLoc m:val="{limloc}"/>'
              f'</m:naryPr>')
    return (f"<m:nary>{naryPr}<m:sub>{lo}</m:sub><m:sup>{hi}</m:sup>"
            f"<m:e>{body}</m:e></m:nary>")


def func(name: str, arg: str) -> str:
    return (f"<m:func><m:fName>{op(name)}</m:fName><m:e>{arg}</m:e></m:func>")


def under(op_name: str, sub_: str) -> str:
    """operator with a subscript limit below, e.g. max_j, argmin."""
    return sub(op(op_name), sub_)


def indicator(inner: str) -> str:
    """1[ inner ]  (upright bold-ish one + brackets)."""
    return op("1") + brack(inner)


def concat(*parts: str) -> str:
    return "".join(parts)


# ---- insertion ------------------------------------------------------------ #
def add_equation(doc, body: str, number: str = None):
    """Insert a centered display equation, optional right-aligned number."""
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.shared import Inches

    p = doc.add_paragraph()
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")
    p._p.append(parse_xml(f'<m:oMath xmlns:m="{M}">{body}</m:oMath>'))
    if number:
        p.add_run("\t(" + str(number) + ")")
    return p


def inline(paragraph, body: str):
    """Append an inline equation to an existing paragraph."""
    from docx.oxml import parse_xml
    paragraph._p.append(parse_xml(f'<m:oMath xmlns:m="{M}">{body}</m:oMath>'))
    return paragraph
